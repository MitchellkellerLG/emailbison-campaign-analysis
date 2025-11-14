#!/usr/bin/env python3
"""
Convert ALL TeachAid Geography files (Email Copy + Research Reports) to DOCX format
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
import shutil

def parse_markdown_to_docx(md_content, doc_title):
    """Convert markdown content to a formatted Word document"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Main title (# heading)
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # H2 heading (##)
        elif line.startswith('## '):
            heading = line[3:].strip()
            doc.add_heading(heading, level=1)

        # H3 heading (###)
        elif line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=2)

        # H4 heading (####)
        elif line.startswith('#### '):
            heading = line[5:].strip()
            doc.add_heading(heading, level=3)

        # Horizontal rule (---)
        elif line.strip() == '---':
            doc.add_paragraph('_' * 80)

        # Bold subject lines
        elif line.startswith('**Subject:**'):
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', ''))
            run.bold = True
            run.font.size = Pt(12)

        # Bullets (•)
        elif line.startswith('• '):
            doc.add_paragraph(line[2:], style='List Bullet')

        # Numbered bullets (1., 2., etc)
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(line[line.index('.')+2:], style='List Number')

        # Wait instructions
        elif line.startswith('[WAIT'):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)

        # Regular paragraphs
        else:
            # Handle bold text (**text**)
            para_text = line
            p = doc.add_paragraph()

            # Split by bold markers
            parts = re.split(r'(\*\*.*?\*\*)', para_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Bold text
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Normal text
                    p.add_run(part)

        i += 1

    return doc

def convert_all_files():
    """Convert all email campaign AND research report markdown files to docx"""

    # Get current directory
    base_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = os.path.join(base_dir, 'Docx_Versions')

    # Ensure output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # List of ALL files to convert
    files = [
        # Email Campaign Copy
        'ALBERTA_Email_Copy_Variants.md',
        'MANITOBA_Email_Copy_Variants.md',
        'QUEBEC_Email_Copy_Variants.md',
        'ONTARIO_Email_Copy_Variants.md',
        'BRITISH_COLUMBIA_Email_Copy_Variants.md',
        'OHIO_Email_Copy_Variants.md',
        'SOUTH_AFRICA_Email_Copy_Variants.md',
        'UK_Email_Copy_Variants.md',

        # Research Reports
        'Alberta_Research_Report.md',
        'Manitoba_Research_Report.md',
        'Quebec_Research_Report.md',
        'Ontario_Research_Report.md',
        'British_Columbia_Research_Report.md',
        'Ohio_Research_Report.md',
        'South_Africa_Research_Report.md',
        'UK_Research_Report.md',

        # Master Summary
        'MASTER_SUMMARY_AND_DELIVERABLES.md'
    ]

    print("Converting ALL TeachAid Geography files to DOCX...\n")
    print("=" * 60)

    converted_count = 0

    for filename in files:
        filepath = os.path.join(base_dir, filename)

        if not os.path.exists(filepath):
            print(f"WARNING: Skipping {filename} (not found)")
            continue

        # Read markdown content
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()

        # Extract document title
        doc_title = filename.replace('.md', '').replace('_', ' ')

        # Convert to docx
        doc = parse_markdown_to_docx(content, doc_title)

        # Save as docx in Docx_Versions folder
        output_filename = filename.replace('.md', '.docx')
        output_path = os.path.join(output_dir, output_filename)
        doc.save(output_path)

        print(f"Created: {output_filename}")
        converted_count += 1

    print("=" * 60)
    print(f"\nAll {converted_count} files converted successfully!")
    print(f"\nLocation: {output_dir}")

    # Move existing docx files to new folder
    print("\nMoving existing DOCX files to Docx_Versions folder...")
    existing_docx = [f for f in os.listdir(base_dir) if f.endswith('.docx') and os.path.isfile(os.path.join(base_dir, f))]

    moved_count = 0
    for docx_file in existing_docx:
        src = os.path.join(base_dir, docx_file)
        dst = os.path.join(output_dir, docx_file)
        # Only move if not already there
        if src != dst:
            try:
                shutil.move(src, dst)
                print(f"Moved: {docx_file}")
                moved_count += 1
            except:
                pass  # Already exists in destination

    if moved_count > 0:
        print(f"\nMoved {moved_count} existing DOCX files")

    print("\nDONE! All DOCX files are now in Docx_Versions folder")

if __name__ == '__main__':
    convert_all_files()
