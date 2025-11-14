#!/usr/bin/env python3
"""
Convert TeachAid Geography Email Campaign Markdown files to DOCX format
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def parse_markdown_to_docx(md_content, geography_name):
    """Convert markdown content to a formatted Word document"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Main title (# heading)
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # H2 heading (##)
        elif line.startswith('## '):
            heading = line[3:].strip()
            doc.add_heading(heading, level=1)

        # H3 heading (###)
        elif line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=2)

        # Horizontal rule (---)
        elif line.strip() == '---':
            doc.add_paragraph('_' * 80)

        # Bold subject lines
        elif line.startswith('**Subject:**'):
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', ''))
            run.bold = True
            run.font.size = Pt(12)

        # Bullets (•)
        elif line.startswith('• '):
            doc.add_paragraph(line[2:], style='List Bullet')

        # Wait instructions
        elif line.startswith('[WAIT'):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)

        # Regular paragraphs
        else:
            # Handle bold text (**text**)
            para_text = line
            p = doc.add_paragraph()

            # Split by bold markers
            parts = re.split(r'(\*\*.*?\*\*)', para_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Bold text
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Normal text
                    p.add_run(part)

        i += 1

    return doc

def convert_all_markdown_files():
    """Convert all email campaign markdown files to docx"""

    # Get current directory
    base_dir = os.path.dirname(os.path.abspath(__file__))

    # List of geography files to convert
    files = [
        'ALBERTA_Email_Copy_Variants.md',
        'MANITOBA_Email_Copy_Variants.md',
        'QUEBEC_Email_Copy_Variants.md',
        'ONTARIO_Email_Copy_Variants.md',
        'BRITISH_COLUMBIA_Email_Copy_Variants.md',
        'OHIO_Email_Copy_Variants.md',
        'SOUTH_AFRICA_Email_Copy_Variants.md',
        'UK_Email_Copy_Variants.md'
    ]

    print("Converting TeachAid Geography Email Campaign files to DOCX...\n")

    for filename in files:
        filepath = os.path.join(base_dir, filename)

        if not os.path.exists(filepath):
            print(f"WARNING: Skipping {filename} (not found)")
            continue

        # Read markdown content
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()

        # Extract geography name
        geography = filename.replace('_Email_Copy_Variants.md', '').replace('_', ' ')

        # Convert to docx
        doc = parse_markdown_to_docx(content, geography)

        # Save as docx
        output_filename = filename.replace('.md', '.docx')
        output_path = os.path.join(base_dir, output_filename)
        doc.save(output_path)

        print(f"Created: {output_filename}")

    print("\nAll files converted successfully!")
    print(f"\nLocation: {base_dir}")

if __name__ == '__main__':
    convert_all_markdown_files()
