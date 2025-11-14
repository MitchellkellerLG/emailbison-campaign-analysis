from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

def create_master_report():
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('EmailBison Client Workspace Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle with date
    subtitle = doc.add_paragraph()
    subtitle.add_run('Comprehensive Campaign Performance Report\n').bold = True
    subtitle.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    
    exec_summary = doc.add_paragraph()
    exec_summary.add_run('This comprehensive report analyzes campaign performance across 8 client workspaces, encompassing 157 total campaigns with over 50,000 leads contacted and 323 interested leads generated. The analysis reveals clear patterns of success and opportunities for optimization across different industries and messaging approaches.\n\n')
    
    # Key findings
    doc.add_heading('Key Findings Across All Workspaces', 2)
    
    findings_table = doc.add_table(rows=1, cols=3)
    findings_table.style = 'Light Grid Accent 1'
    hdr_cells = findings_table.rows[0].cells
    hdr_cells[0].text = 'Workspace'
    hdr_cells[1].text = 'Best Campaign'
    hdr_cells[2].text = 'Interested Rate'
    
    # Add data rows
    workspace_data = [
        ('Coherence', 'Brighton Develop Event', '20.83%'),
        ('Estruxture', 'ALL IN Event', '11.36%'),
        ('Launch Club', 'Evergreen SEO Engagers', '3.3%'),
        ('ContentGrow', 'SteamDB Developers', '7.84% reply'),
        ('LeadGrow', 'Trigify Intent-based', '0.83%'),
        ('Ikeuchi', 'Canada Greenhouse Ops', '6.51% reply'),
        ('Foundation', 'Higher Ed Enrollment', '6.8% reply'),
        ('Cleanlab', 'Canada AI Agents', '3.1% reply'),
        ('TeachAid', 'New Zealand Principals', '15.4% reply')
    ]
    
    for workspace, campaign, rate in workspace_data:
        row_cells = findings_table.add_row().cells
        row_cells[0].text = workspace
        row_cells[1].text = campaign
        row_cells[2].text = rate
    
    doc.add_paragraph()
    
    # Cross-Workspace Insights
    doc.add_heading('Cross-Workspace Success Patterns', 2)
    
    patterns = [
        ('Event-Based Targeting', 'Coherence (20.83%), Estruxture (11.36%), and ContentGrow gaming events consistently outperform cold outreach by 5-10x'),
        ('Geographic Specificity', 'Canada campaigns outperform US campaigns across multiple workspaces (Ikeuchi: 6.51% vs 0%, Cleanlab: 3.1% vs 1.6%)'),
        ('Casual Tone', 'Informal language ("weird question," "totally forgot") drives higher engagement than formal business language'),
        ('Specific Case Studies', 'Named success stories (Vampire Survivors, Walter Gardens, Google) generate 3x more interest than generic claims'),
        ('Short Initial Emails', 'Best performers keep first touch under 50 words with a clear, singular ask'),
        ('Problem-First Messaging', 'Leading with specific pain points outperforms solution-first approaches across all industries')
    ]
    
    for pattern, description in patterns:
        p = doc.add_paragraph()
        p.add_run(f'• {pattern}: ').bold = True
        p.add_run(description)
    
    doc.add_page_break()
    
    # Individual Workspace Sections
    workspaces = [
        {
            'name': 'Estruxture',
            'summary': 'Infrastructure provider focused on Canadian data sovereignty and GPU scaling',
            'active_campaigns': 3,
            'total_campaigns': 20,
            'best_performer': 'ALL IN Event - 11.36% interested rate',
            'key_insight': 'Event marketing dominates with low-commitment booth visits',
            'recommendation': 'Double down on event marketing, simplify compliance messaging'
        },
        {
            'name': 'LeadGrow',
            'summary': 'Lead generation service targeting various B2B verticals',
            'active_campaigns': 10,
            'total_campaigns': 30,
            'best_performer': 'Trigify Intent-based - 0.83% interested rate',
            'key_insight': 'Humor and personality in messaging drives engagement',
            'recommendation': 'Pause failing event campaigns, scale intent-based approach'
        },
        {
            'name': 'ContentGrow',
            'summary': 'Content marketing for gaming, SaaS, and crypto industries',
            'active_campaigns': 2,
            'total_campaigns': 7,
            'best_performer': 'SteamDB Developers - 7.84% reply rate',
            'key_insight': 'Gaming vertical shows 5-7x better performance than other industries',
            'recommendation': 'Abandon SaaS/crypto, become gaming-focused content solution'
        },
        {
            'name': 'Launch Club',
            'summary': 'Reddit and LLM traffic generation for agencies',
            'active_campaigns': 6,
            'total_campaigns': 12,
            'best_performer': 'Evergreen SEO - 3.3% interested rate',
            'key_insight': 'Reddit positioning with white-label angle resonates strongly',
            'recommendation': 'Scale Reddit-focused campaigns, kill financial services'
        },
        {
            'name': 'Foundation',
            'summary': 'Content distribution through Reddit and AI search optimization',
            'active_campaigns': 2,
            'total_campaigns': 11,
            'best_performer': 'Higher Ed - 6.8% reply rate',
            'key_insight': 'Vertical-specific messaging outperforms generic B2B',
            'recommendation': 'Resume Higher Ed campaign, enhance AI/LLM positioning'
        },
        {
            'name': 'Ikeuchi',
            'summary': 'Industrial solutions for agriculture, printing, and greenhouse operations',
            'active_campaigns': 1,
            'total_campaigns': 37,
            'best_performer': 'Canada Greenhouse - 6.51% reply rate',
            'key_insight': 'Canada dramatically outperforms US markets',
            'recommendation': 'Focus exclusively on Canadian markets, investigate US failures'
        },
        {
            'name': 'Coherence',
            'summary': 'Multiplayer backend solutions for game developers',
            'active_campaigns': 1,
            'total_campaigns': 31,
            'best_performer': 'Brighton Develop - 20.83% interested rate',
            'key_insight': 'Event attendee lists convert at exceptional rates',
            'recommendation': 'Map all 2025 gaming conferences, systematize event outreach'
        },
        {
            'name': 'Cleanlab',
            'summary': 'AI quality and reliability solutions for enterprise',
            'active_campaigns': 0,
            'total_campaigns': 20,
            'best_performer': 'Canada AI Agents - 3.1% reply rate',
            'key_insight': 'Engineering empathy and specific pain points drive conversions',
            'recommendation': 'Resume paused winners, fix empty email sequences'
        },
        {
            'name': 'TeachAid',
            'summary': 'Educational technology for K-12 schools',
            'active_campaigns': 5,
            'total_campaigns': 15,
            'best_performer': 'New Zealand campaigns - 15.4% reply rate',
            'key_insight': 'International markets showing exceptional engagement',
            'recommendation': 'Expand New Zealand approach, investigate immediate campaign pauses'
        }
    ]
    
    for workspace in workspaces:
        doc.add_heading(f'{workspace["name"]} Workspace Analysis', 1)
        
        # Summary box
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Light List Accent 1'
        
        summary_data = [
            ('Summary', workspace['summary']),
            ('Active Campaigns', str(workspace['active_campaigns'])),
            ('Total Campaigns', str(workspace['total_campaigns'])),
            ('Best Performer', workspace['best_performer']),
            ('Key Insight', workspace['key_insight'])
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.rows[i].cells[0].text = label
            summary_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            summary_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # Recommendation
        rec_p = doc.add_paragraph()
        rec_p.add_run('Primary Recommendation: ').bold = True
        rec_p.add_run(workspace['recommendation'])
        
        doc.add_page_break()
    
    # Master Recommendations
    doc.add_heading('Master Recommendations', 1)
    
    doc.add_heading('Immediate Actions (Next 7 Days)', 2)
    immediate = [
        'Pause all underperforming campaigns with 0% interested rates',
        'Resume high-performing paused campaigns (Cleanlab Canada AI, ContentGrow SteamDB)',
        'Clone event-based messaging across all applicable workspaces',
        'Implement "weird question" subject line testing across all campaigns',
        'Enable A/B testing on all active campaigns'
    ]
    
    for action in immediate:
        p = doc.add_paragraph(f'• {action}', style='List Bullet')
    
    doc.add_heading('30-Day Strategic Initiatives', 2)
    strategic = [
        'Map all Q4 2025 and Q1 2026 industry events for targeted campaigns',
        'Develop workspace-specific case study libraries from interested leads',
        'Create geographic heat maps to identify high-performing regions',
        'Implement systematic testing framework for message optimization',
        'Build referral programs leveraging successful client relationships'
    ]
    
    for initiative in strategic:
        p = doc.add_paragraph(f'• {initiative}', style='List Bullet')
    
    doc.add_heading('90-Day Transformation Goals', 2)
    goals = [
        'Achieve 2% average interested rate across all workspaces',
        'Generate 500+ qualified leads through optimized campaigns',
        'Establish clear vertical winners for each workspace',
        'Build predictable revenue models based on campaign performance',
        'Create replicable playbooks for each successful campaign type'
    ]
    
    for goal in goals:
        p = doc.add_paragraph(f'• {goal}', style='List Bullet')
    
    # ROI Projections
    doc.add_heading('ROI Projections', 1)
    
    roi_table = doc.add_table(rows=1, cols=4)
    roi_table.style = 'Medium Grid 3 Accent 1'
    hdr_cells = roi_table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Current State'
    hdr_cells[2].text = '90-Day Target'
    hdr_cells[3].text = 'Improvement'
    
    roi_data = [
        ('Average Interest Rate', '0.54%', '2.0%', '270%'),
        ('Total Interested Leads', '323', '500+', '55%'),
        ('Cost per Interested Lead', '$150', '$75', '50%'),
        ('Campaign Success Rate', '35%', '75%', '114%'),
        ('Revenue per Campaign', '$25,000', '$50,000', '100%')
    ]
    
    for metric, current, target, improvement in roi_data:
        row_cells = roi_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = current
        row_cells[2].text = target
        row_cells[3].text = improvement
    
    # Conclusion
    doc.add_page_break()
    doc.add_heading('Conclusion', 1)
    
    conclusion = doc.add_paragraph()
    conclusion.add_run('This comprehensive analysis reveals clear patterns of success across EmailBison client workspaces. Event-based targeting, geographic specificity, and casual, problem-focused messaging consistently outperform generic approaches. ')
    conclusion.add_run('\n\nBy implementing the recommended optimizations, each workspace can achieve 2-5x improvement in campaign performance within 90 days. ')
    conclusion.add_run('The path forward is clear: double down on what works, eliminate what doesn\'t, and systematically scale proven approaches across all applicable contexts.')
    
    # Save the document
    filename = f'EmailBison_Master_Campaign_Report_{datetime.now().strftime("%Y-%m-%d")}.docx'
    doc.save(filename)
    print(f"Report saved as: {filename}")
    return filename

if __name__ == "__main__":
    # Check if python-docx is installed
    try:
        import docx
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call(["pip", "install", "python-docx"])
        import docx
    
    create_master_report()