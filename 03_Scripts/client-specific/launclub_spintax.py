from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Launclub Email Sequences - Heavily Spintaxed (3 Variants)', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Josh Braun/Will Allred Style with Anti-Fingerprinting Spintax')
doc.add_paragraph('Variables: {FIRST_NAME}, {SENDER_EMAIL_SIGNATURE} only')
doc.add_paragraph()

# Campaign 1: CEO/Founder/People Leaders
doc.add_heading('CAMPAIGN 1: New Hires - CEO, Founder, People Leaders', 1)

# Step 1
doc.add_heading('STEP 1 (3 Days Wait)', 2)

step1_v1 = {
    'subject': '{reddit mayhem|search traffic shift|reddit taking over}',
    'body': '''{Hey|Hi} {FIRST_NAME}, {saw|noticed|watching how} Reddit {basically hijacked|absorbed|pulled in} {20%|a fifth|one-fifth} of {Google search|search traffic|search volume} this year.

{Quick question|Wondering|Quick ask} - would it {help|be useful|make sense} if I {sent over|fired over|shared} a {breakdown|report|audit} of where your Reddit opportunities are and what to do about them?

{No pitch.|Free, obviously.|No strings.} {Just figured it might be useful.|Think it'd help.|Figure it could be valuable.}

{SENDER_EMAIL_SIGNATURE}

P.S. We {built|handle|created} Reddit {strategies|programs|initiatives} for {Blizzard, Afterpay, and one of the Big 3 CRMs|brands like Blizzard and Afterpay|orgs like Blizzard, Afterpay, and a Top 3 CRM}.'''
}

step1_v2 = {
    'subject': '{your reddit blind spot|20% search moved|reddit eating volume}',
    'body': '''{Hi|Hey} {FIRST_NAME},

Reddit now {owns|controls|drives} {20%|one-fifth} of search. {Most companies aren't adapting.|Wild.|Not many are paying attention.}

I can {send you|fire over|share} a {free|complimentary|no-cost} {audit|report|breakdown} showing your {specific|current|exact} opportunities on Reddit - where to post, what to say, how to convert.

{Interested?|Worth a look?|Make sense?}

{SENDER_EMAIL_SIGNATURE}

P.S. We {handle|manage|run} Reddit for {Top 3 CRM and gaming companies like Blizzard|some of the largest B2B SaaS|brands like Blizzard and Afterpay} - usually {uncovers|surfaces|reveals} {10-30k|10,000-30,000|ten to thirty thousand} monthly visitors.'''
}

step1_v3 = {
    'subject': '{reddits search takeover|untapped reddit angle|reddit opportunities}',
    'body': '''{Hey|Hi} {FIRST_NAME}, {wild idea|crazy thought|bold idea} - {want me to|would it be alright if I|interested if I} {send you|fire over|shoot you} a {free|no-cost|complimentary} {audit|report|breakdown} that {outlines|covers|maps out} your current Reddit opportunities and how to capitalize?

{Could help you hit the ground running.|Think it'd be useful.|Figure it might help.}

{Let me know, thanks.|Lmk.|Worth a look?}

{SENDER_EMAIL_SIGNATURE}

P.S. {It typically|Usually|Generally} {uncovers|surfaces|reveals} {10-30k|10,000-30,000} in monthly {product-led|high-intent|buying-intent} search volume.'''
}

doc.add_heading('VARIANT 1', 3)
doc.add_paragraph('Subject: ' + step1_v1['subject'])
doc.add_paragraph(step1_v1['body'])

doc.add_heading('VARIANT 2', 3)
doc.add_paragraph('Subject: ' + step1_v2['subject'])
doc.add_paragraph(step1_v2['body'])

doc.add_heading('VARIANT 3', 3)
doc.add_paragraph('Subject: ' + step1_v3['subject'])
doc.add_paragraph(step1_v3['body'])

# Step 2
doc.add_heading('STEP 2 (Wait 3 Days After Step 1)', 2)

step2_v1 = {
    'subject': 'Re: {reddit mayhem|search traffic shift|reddit taking over}',
    'body': '''{Probably should've been clearer|Realized I didn't explain this clearly|Could've been more specific} {FIRST_NAME}.

The report {shows|maps out|breaks down} Reddit traffic opportunities based on {searches|keywords|rankings} you already {rank for|have rankings on|show up for}. {Which subreddits to focus on.|You'll see exact subreddits to target.|Target subreddits identified.} {What posts to create.|Types of posts that convert.|Content to publish.} How to turn Reddit {browsers|visitors|traffic} into customers {(and get picked up by LLMs while you're at it)|(plus show up in LLM responses)|(and capture LLM mentions)}.

If {it's helpful|you think it'd be useful|this sounds good}, {let me know|give me a shout|lmk} and I'll {get it over to you|put it together|send it over} ASAP.

{SENDER_EMAIL_SIGNATURE}'''
}

step2_v2 = {
    'subject': 'Re: {your reddit blind spot|20% search moved|reddit eating volume}',
    'body': '''{Thinking about it|On reflection|Looking back}, I {could've done a better job explaining this|should've clarified better|didn't explain clearly enough} {FIRST_NAME}.

{Basically,|Here's the thing -|In short,} the {audit|report|breakdown} {breaks down|analyzes|maps out} Reddit {traffic|visitor|growth} opportunities {tied to|based on|linked to} your current rankings.

{Shows you|You'll see|You get}:
- {Target|Priority|Key} subreddits
- Post types that {drive conversions|convert|generate customers}
- How to {get traction in|show up in|capture} LLM results

{If you think it'd be useful|Useful?|Sound helpful?} {Give me a shout|Let me know|Lmk} and I'll {get it ready|throw it together|put it together}.

{SENDER_EMAIL_SIGNATURE}'''
}

step2_v3 = {
    'subject': 'Re: {reddits search takeover|untapped reddit angle|reddit opportunities}',
    'body': '''{Let me clarify|Should've been clearer|Didn't explain this well} {FIRST_NAME}.

This {report|breakdown|analysis} {analyzes|examines|maps out} Reddit opportunities based on the searches you already {rank on|show up for|have rankings for}. It {tells you|shows you|outlines} which subreddits to post in, what content to create, and how to convert Reddit traffic.

{If it sounds helpful|Worth having?|Make sense?} {Lmk|Let me know|Give me a shout} and I'll {put it together fast|send it over ASAP|get it ready quickly}.

{SENDER_EMAIL_SIGNATURE}'''
}

doc.add_heading('VARIANT 1', 3)
doc.add_paragraph('Subject: ' + step2_v1['subject'])
doc.add_paragraph(step2_v1['body'])

doc.add_heading('VARIANT 2', 3)
doc.add_paragraph('Subject: ' + step2_v2['subject'])
doc.add_paragraph(step2_v2['body'])

doc.add_heading('VARIANT 3', 3)
doc.add_paragraph('Subject: ' + step2_v3['subject'])
doc.add_paragraph(step2_v3['body'])

# Step 3
doc.add_heading('STEP 3 (Wait 1 Day After Step 2)', 2)

step3_v1 = {
    'subject': 'Re: {reddit mayhem|search traffic shift|reddit taking over}',
    'body': '''{FIRST_NAME}, we've {completed|done|run} about 10 of these {this week|over the past few days|recently}. Results have been {solid|promising|impressive}.

You {walk away with|get|receive}:
- 10 priority subreddits to {post in|focus on|target}
- {Clear path|Strategy|Plan} to {10-30k monthly visitors|10,000-30,000 monthly reach|capture 10-30k audience growth} from Reddit
- Exactly what to post{, where, and when|and where| (what, where, how)}

{Not sure if you're the right contact though.|I'm not certain if you're the best person for this.|Unsure if you're the one who should see this.} Should I be {talking to|connecting with|speaking with} someone else on your team?

{Either way, let me know.|Appreciate your guidance either way.|Let me know - thanks.}

{SENDER_EMAIL_SIGNATURE}'''
}

step3_v2 = {
    'subject': 'Re: {your reddit blind spot|20% search moved|reddit eating volume}',
    'body': '''{Hey|Hi} {FIRST_NAME}, {we've done roughly|completed about|wrapped up} 10 of these {in the last week|recently|this week}. {Results have been quite good.|Outcomes have been strong.|Pretty impressive results.}

{You'll walk away with|What you get|Here's what you receive}:
- 10 priority subreddits
- How to {generate|achieve|capture} {10-30k monthly visitors|10,000-30,000 monthly audience growth|ten to thirty thousand visitors monthly} from Reddit
- {Exactly what to post and where|Detailed posting strategy|Exact content strategy}

{Not certain if this should go to you or someone else on your team though.|Unsure if this should go to you though.|I'm not sure if you're the right person.} {Thoughts?|Would someone else be better?|Should I connect with another teammate?}

{Thanks either way.|Let me know, thanks.|Appreciate it.}

{SENDER_EMAIL_SIGNATURE}'''
}

step3_v3 = {
    'subject': 'Re: {reddits search takeover|untapped reddit angle|reddit opportunities}',
    'body': '''{FIRST_NAME} - we {ran|completed|finished} about 10 of these {over the past week|recently|this week}. {Results look strong.|Pretty good outcomes.|Solid results.}

{What you walk away with|You get|Here's what you receive}:
- 10 priority subreddits
- Path to {10-30k monthly visitors|capture 10,000-30,000 monthly reach|generate ten to thirty thousand monthly visitors} from Reddit
- {Exact posts to make, and where|What to post, where, and when|Detailed content and posting strategy}

{Not sure if you're the right person though|I'm not certain if this should go to you|Unsure if you're the best contact} - {is there someone else on your team I should send it to?|would someone else on your team be better?|should I be talking to another teammate?}

{Let me know.|Either way, thanks.|Appreciate your help.}

{SENDER_EMAIL_SIGNATURE}'''
}

doc.add_heading('VARIANT 1', 3)
doc.add_paragraph('Subject: ' + step3_v1['subject'])
doc.add_paragraph(step3_v1['body'])

doc.add_heading('VARIANT 2', 3)
doc.add_paragraph('Subject: ' + step3_v2['subject'])
doc.add_paragraph(step3_v2['body'])

doc.add_heading('VARIANT 3', 3)
doc.add_paragraph('Subject: ' + step3_v3['subject'])
doc.add_paragraph(step3_v3['body'])

# CAMPAIGN 2
doc.add_page_break()
doc.add_heading('CAMPAIGN 2: New Hires - Marketing Leadership', 1)

# Step 1
doc.add_heading('STEP 1 (3 Days Wait)', 2)

step1_mk_v1 = {
    'subject': '{reddits overlooked value|reddit opportunities|new role reddit audit}',
    'body': '''{Hey|Hi} {FIRST_NAME}, {saw|noticed|just saw} you joined your {company|org|new company} as {a marketing leader|in a marketing role|in marketing leadership} {recently|pretty recently|not too long ago}.

Would it {help|be useful|be valuable} if I {fire over|send you|shoot across} a {no-cost|complimentary|free} {audit|report|deck} {outlining|mapping out|highlighting} your current {opportunities|potential wins|growth paths} on Reddit so you can hit the ground running?

{SENDER_EMAIL_SIGNATURE}'''
}

step1_mk_v2 = {
    'subject': '{quick reddit win|reddit growth path|ramping up quickly}',
    'body': '''{FIRST_NAME}, I {figure|imagine|bet} ramping up {quickly|smoothly|fast} in your new {marketing role|role|position} is {top of mind|a big priority|front and center} {right now|these days|atm}.

{Quick question|I wanted to ask|Just wondering} - would it be {alright|okay|cool} if I {sent over|shared|fired across} a {complimentary|no-cost|value-packed} report {covering|highlighting|mapping out} your current {opportunities|angles|potential wins} on Reddit?

{SENDER_EMAIL_SIGNATURE}'''
}

step1_mk_v3 = {
    'subject': '{reddit opportunities|reddit growth report|new role support}',
    'body': '''{Hey|Hi} {FIRST_NAME}, would it be {alright|okay|cool} if I {fired over|sent you|shot you} a {complimentary|no-cost|on-the-house} report {covering|outlining|breaking down} your current Reddit opportunities so you can {hit the ground running|ramp up quickly|ramp up fast} in your {fresh|new} role?

You'll see {where to post, what to post, and how to get high-intent customers|where to post, what content works, how to convert customers|target subreddits, content strategy, customer conversion tactics}.

{SENDER_EMAIL_SIGNATURE}'''
}

doc.add_heading('VARIANT 1', 3)
doc.add_paragraph('Subject: ' + step1_mk_v1['subject'])
doc.add_paragraph(step1_mk_v1['body'])

doc.add_heading('VARIANT 2', 3)
doc.add_paragraph('Subject: ' + step1_mk_v2['subject'])
doc.add_paragraph(step1_mk_v2['body'])

doc.add_heading('VARIANT 3', 3)
doc.add_paragraph('Subject: ' + step1_mk_v3['subject'])
doc.add_paragraph(step1_mk_v3['body'])

# Step 2
doc.add_heading('STEP 2 (Wait 3 Days After Step 1)', 2)

step2_mk_v1 = {
    'subject': 'Re: {reddits overlooked value|reddit opportunities|new role reddit audit}',
    'body': '''{Reason I ask|The reason I'm reaching out|Let me explain why this matters} {FIRST_NAME} is because {since|after} the {Google + Reddit|Reddit-Google|Reddit and Google} partnership {this year|earlier this year}, Reddit now drives {roughly|about|nearly} 90% of high-intent searches.

This report will {outline|map out|highlight} where your customers are {spending time|active|showing up} and how to {engage with them|connect with them|interact with them} so this becomes {a measurable growth channel|real business impact|more than vanity metrics}.

{For example|Case in point|To illustrate}, when we {worked with|partnered with|helped} a personal finance app during the Mint shutdown{:|, they saw:|}

{6.2x growth|6.2x increase|6.2x lift} in organic sessions
{4.9x increase|4.9x growth|4.9x lift} in installs from Reddit
{5.1x lift|5.1x increase|5.1x growth} from AI search mentions
{37% reduction|37% drop|37% decrease} in support tickets

{Nothing expected in return|Just sharing in case it's useful|Happy to pass this along} - would you like us to {share|provide|put together} a report that {pinpoints|surfaces|maps out} {10-30k|10,000-30,000|ten to thirty thousand} high-priority potential customers?

{SENDER_EMAIL_SIGNATURE}

P.S. We {handle|manage|run} Reddit for some of the {largest B2B SaaS in the world|biggest B2B SaaS globally|world's largest B2B SaaS}.'''
}

step2_mk_v2 = {
    'subject': 'Re: {quick reddit win|reddit growth path|ramping up quickly}',
    'body': '''{Here's why I'm asking|The reason I'm mentioning this|Let me clarify why} {FIRST_NAME}. {After|Since} the {Google + Reddit|Reddit-Google} partnership{,| earlier this year,} Reddit {drives about|now drives roughly|drives nearly} 90% of high-intent searches.

{This|The} report will {show you|outline|highlight} where your customers are {spending time|active|showing up} and {how to engage them|exactly how to connect|how to interact with them} to {turn this into real business impact|make this a measurable growth channel|drive more than vanity metrics}.

{For reference|As an example|To illustrate}, we {worked with|partnered with|helped} a personal finance app during {the Mint shutdown|Mint's shutdown}. {They experienced|Results were|They saw}:

{6.2x organic session growth|6.2x growth in organic sessions|6.2x increase in sessions}
{4.9x Reddit install lift|4.9x increase in Reddit installs|4.9x Reddit-driven installs}
{5.1x increase in AI search mentions|5.1x AI search lift|5.1x lift from AI search}
{37% support ticket reduction|37% drop in support tickets|37% decrease in tickets}

{Nothing expected|Just sharing in case it helps|Happy to share this} - {want us to|would you like us to|interested in having us} {send over|share|provide} a report {identifying|surfacing|mapping out} {10-30k|10,000-30,000} high-priority customers?

{SENDER_EMAIL_SIGNATURE}

P.S. We {run|handle|manage} Reddit for some of the {biggest B2B SaaS in the world|largest B2B SaaS globally|world's largest B2B SaaS}.'''
}

step2_mk_v3 = {
    'subject': 'Re: {reddit opportunities|reddit growth report|new role support}',
    'body': '''The reason {I'm asking|I'm reaching out|this matters} {FIRST_NAME} is {that since|because after|that following} the {Reddit-Google|Google + Reddit} partnership, Reddit now drives {roughly|about|nearly} 90% of high-intent searches.

The {report|breakdown|analysis} will {map out|outline|show} where your customers are {active|spending time|showing up} and how to {connect with them|engage them|interact} so {you get real business impact|this becomes a measurable growth channel|this drives more than vanity metrics}.

{Example|Case in point|For reference}: we {partnered with|worked with|helped} a personal finance app during {Mint's shutdown|the Mint shutdown}. Results:

{6.2x organic sessions|6.2x session growth|6.2x organic session increase}
{4.9x Reddit-driven installs|4.9x installs from Reddit|4.9x Reddit install lift}
{5.1x AI search lift|5.1x increase in AI search mentions|5.1x lift from AI search}
{37% drop in support tickets|37% support ticket reduction|37% decrease in tickets}

{Just sharing in case it helps|Nothing expected|Happy to pass this along} - want us to {share|provide|send over} a report that {surfaces|identifies|pinpoints} {10-30k|10,000-30,000} high-priority potential customers?

{SENDER_EMAIL_SIGNATURE}

P.S. We {handle|manage|run} Reddit for some of the {largest B2B SaaS globally|biggest B2B SaaS in the world|world's largest B2B SaaS companies}.'''
}

doc.add_heading('VARIANT 1', 3)
doc.add_paragraph('Subject: ' + step2_mk_v1['subject'])
doc.add_paragraph(step2_mk_v1['body'])

doc.add_heading('VARIANT 2', 3)
doc.add_paragraph('Subject: ' + step2_mk_v2['subject'])
doc.add_paragraph(step2_mk_v2['body'])

doc.add_heading('VARIANT 3', 3)
doc.add_paragraph('Subject: ' + step2_mk_v3['subject'])
doc.add_paragraph(step2_mk_v3['body'])

# Step 3
doc.add_heading('STEP 3 (Wait 1 Day After Step 2)', 2)

step3_mk_v1 = {
    'subject': 'Re: {reddits overlooked value|reddit opportunities|new role reddit audit}',
    'body': '''{Looking back|Thinking about it|On reflection}, maybe I've {missed the mark|reached out to the wrong person|got this a little off} {FIRST_NAME}. Are you the right person to {review|get|see} a report that {uncovers|highlights|surfaces} {overlooked|untapped|under-the-radar} customers on Reddit and {outlines|details|shows} exactly how to {capture|engage|connect with} them?

{Let me know|Please let me know|Happy to hear} if I should {speak with|connect with|reach out to} someone else.

{Hope you crush it in the new role.|Wishing you the best in the new role.|Best of luck as you ramp up.}

{SENDER_EMAIL_SIGNATURE}'''
}

step3_mk_v2 = {
    'subject': 'Re: {quick reddit win|reddit growth path|ramping up quickly}',
    'body': '''{Thinking about it|Looking back|On reflection}, {maybe I've reached out to the wrong person|I might have missed the mark|maybe I've got this off} {FIRST_NAME}. Are you the right {contact|person} to {receive|get|see} a report that {identifies|uncovers|highlights} {overlooked|untapped} Reddit customers and {outlines|shows|details} how to {capture|engage with|connect with} them?

{Let me know|Please let me know|Happy to hear} if I should be {talking to|connecting with|speaking with} someone else.

{Hope the new role goes well.|Wishing you success in your new role.|Best of luck in the new role.}

{SENDER_EMAIL_SIGNATURE}'''
}

step3_mk_v3 = {
    'subject': 'Re: {reddit opportunities|reddit growth report|new role support}',
    'body': '''{On reflection|Thinking about it|Looking back}, maybe I've {reached out to the wrong person|missed the mark|got this a little off} {FIRST_NAME}. Are you the right contact to get a report {uncovering|surfacing|identifying} {untapped|overlooked|under-the-radar} Reddit customers and showing exactly how to {engage them|connect with them|capture them}?

{Please let me know|Let me know|Happy to hear} if {there's someone else I should connect with|I should reach out to someone else|I should be talking to another teammate}.

{Wishing you success in your new role.|Best of luck as you ramp up.|Hope you crush it in the new role.}

{SENDER_EMAIL_SIGNATURE}'''
}

doc.add_heading('VARIANT 1', 3)
doc.add_paragraph('Subject: ' + step3_mk_v1['subject'])
doc.add_paragraph(step3_mk_v1['body'])

doc.add_heading('VARIANT 2', 3)
doc.add_paragraph('Subject: ' + step3_mk_v2['subject'])
doc.add_paragraph(step3_mk_v2['body'])

doc.add_heading('VARIANT 3', 3)
doc.add_paragraph('Subject: ' + step3_mk_v3['subject'])
doc.add_paragraph(step3_mk_v3['body'])

# Add notes section
doc.add_page_break()
doc.add_heading('IMPLEMENTATION NOTES', 1)

notes = [
    ('Variables Used', 'Only {FIRST_NAME} and {SENDER_EMAIL_SIGNATURE} - no custom company, title, or hook variables to reduce fingerprinting patterns'),
    ('Spintax Strategy', 'Heavy spintax on: greetings, action verbs, value propositions, transitions, social proof elements, and CTAs'),
    ('Spam Avoidance', 'Avoided: excessive caps, multiple exclamation marks, "free" overuse, urgency language, clickbait phrases, promotional language'),
    ('Josh Braun Style', 'Short sentences, conversational tone, pattern interrupts, curiosity gaps, no corporate speak'),
    ('Will Allred Style', 'Direct value propositions, social proof in P.S., casual sign-offs, question-based CTAs'),
    ('Thread Continuity', 'Step 2 and 3 use "Re:" to maintain email thread (same subject line spintax as Step 1)'),
    ('Rotation Strategy', 'Rotate between 3 variants randomly per send to avoid ESP pattern detection'),
    ('Wait Times', 'Step 1: Send immediately | Step 2: Wait 3 days | Step 3: Wait 1 day after Step 2')
]

for note_title, note_content in notes:
    p = doc.add_paragraph()
    run = p.add_run(note_title + ': ')
    run.bold = True
    p.add_run(note_content)

# Save
doc.save(r'c:\Users\mitch\Desktop\Claude Code Projects\Launclub_Spintax_3_Variants.docx')
print("Spintax document created successfully!")
