from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Launclub Email Sequence Variants - Anti-Fingerprinting', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Created using Josh Braun/Will Allred copywriting style')
doc.add_paragraph('5 variants per message to combat ESP fingerprinting')
doc.add_paragraph()

# Campaign 1: CEO/Founder/People Leaders
doc.add_heading('CAMPAIGN 1: New Hires - CEO, Founder, People Leaders', 1)

# Step 1 variants
messages_step1_ceo = [
    {
        'title': 'STEP 1 - Message Variant 1',
        'subject': 'reddit mayhem',
        'body': '''Hey {FIRST_NAME}, saw Reddit basically hijacked 20% of Google search this year.

Quick question - would it help if I sent over a breakdown of where your Reddit opportunities are and what to do about them?

No pitch. Just figured it might be useful.

{SENDER_EMAIL_SIGNATURE}

P.S. We built Reddit strategies for Blizzard, Afterpay, and one of the Big 3 CRMs.'''
    },
    {
        'title': 'STEP 1 - Message Variant 2',
        'subject': 'reddit eating search volume',
        'body': '''{FIRST_NAME} - watching how Reddit absorbed a fifth of search traffic this year?

Wondering if you'd find value in a report showing exactly where {COMPANY} could be showing up on Reddit (and how to actually convert those visitors).

Worth a look?

{SENDER_EMAIL_SIGNATURE}

P.S. We've done this for brands like Blizzard and Afterpay - usually uncovers 10-30k monthly visitors.'''
    },
    {
        'title': 'STEP 1 - Message Variant 3',
        'subject': 'your reddit blind spot',
        'body': '''Hi {FIRST_NAME},

Reddit now owns 20% of search. Most companies aren't adapting.

I can send you a free audit showing your specific opportunities on Reddit - where to post, what to say, how to convert.

Interested?

{SENDER_EMAIL_SIGNATURE}

P.S. We handle Reddit for Top 3 CRM and gaming companies like Blizzard.'''
    },
    {
        'title': 'STEP 1 - Message Variant 4',
        'subject': '20% of search moved to reddit',
        'body': '''{FIRST_NAME},

Not sure if this is on your radar, but Reddit pulled in 20% of search volume this year.

Would a report outlining {COMPANY}'s current Reddit opportunities be useful? Shows you exactly which subreddits to target and what content converts.

Free, obviously.

{SENDER_EMAIL_SIGNATURE}

P.S. Built Reddit programs for Afterpay, Blizzard, and a Top 3 CRM.'''
    },
    {
        'title': 'STEP 1 - Message Variant 5',
        'subject': 'reddits search takeover',
        'body': '''Hey {FIRST_NAME},

Reddit absorbed 20% of search this year. Wild.

Quick ask - want me to fire over a breakdown of your Reddit opportunities? It'll show you which subreddits matter and exactly what to post to generate conversions.

Think it'd help.

{SENDER_EMAIL_SIGNATURE}

P.S. We've built Reddit strategies for orgs like Blizzard and Afterpay.'''
    }
]

for msg in messages_step1_ceo:
    doc.add_heading(msg['title'], 2)
    doc.add_paragraph('Subject: ' + msg['subject'])
    doc.add_paragraph(msg['body'])

# Step 2 variants
messages_step2_ceo = [
    {
        'title': 'STEP 2 - Message Variant 1',
        'subject': 'Re: reddit mayhem',
        'body': '''Probably should've been clearer {FIRST_NAME}.

The report shows Reddit traffic opportunities based on searches you already rank for. Which subreddits to focus on. What posts to create. How to turn Reddit browsers into customers (and get picked up by LLMs while you're at it).

If it's helpful, let me know and I'll get it over to you ASAP so {HOOK} can hit the ground running.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 2 - Message Variant 2',
        'subject': 'Re: reddit eating search volume',
        'body': '''{FIRST_NAME} - realized I didn't explain this clearly.

Basically, the report maps out Reddit opportunities based on keywords you currently rank for. You'll see:
- Exact subreddits to target
- Types of posts that convert
- How to capture both Reddit traffic and LLM mentions

Useful for {HOOK} to ramp up quickly in the new role. Want me to throw it together?

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 2 - Message Variant 3',
        'subject': 'Re: your reddit blind spot',
        'body': '''Thinking about it, I could've done a better job explaining this {FIRST_NAME}.

The audit breaks down Reddit traffic opportunities tied to your current rankings. Shows you target subreddits, post types that drive conversions, and how to get traction in LLM results.

If you think it'd be useful for {HOOK} in the new role, give me a shout and I'll get it ready.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 2 - Message Variant 4',
        'subject': 'Re: 20% of search moved to reddit',
        'body': '''On reflection, let me clarify {FIRST_NAME}.

This report analyzes Reddit opportunities based on the searches you already rank on. It tells you which subreddits to post in, what content to create, and how to convert Reddit traffic (plus show up in LLM responses).

If it sounds helpful for {HOOK} to hit the ground running, lmk and I'll put it together fast.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 2 - Message Variant 5',
        'subject': 'Re: reddits search takeover',
        'body': '''Could've been more specific {FIRST_NAME}.

The breakdown shows Reddit opportunities based on keywords you currently rank for. You get: target subreddits, content to publish, conversion tactics, and LLM visibility strategies.

Worth having for {HOOK} to ramp up quickly? Let me know and I'll send it over ASAP.

{SENDER_EMAIL_SIGNATURE}'''
    }
]

for msg in messages_step2_ceo:
    doc.add_heading(msg['title'], 2)
    doc.add_paragraph('Subject: ' + msg['subject'])
    doc.add_paragraph(msg['body'])

# Step 3 variants
messages_step3_ceo = [
    {
        'title': 'STEP 3 - Message Variant 1',
        'subject': 'Re: reddit mayhem',
        'body': '''{FIRST_NAME}, we've completed about 10 of these this week. Results have been solid.

You walk away with:
- 10 priority subreddits to post in
- Clear path to 10-30k monthly visitors from Reddit
- Exactly what to post, where, and when

Not sure if you're the right contact though. Should I be talking to someone else on your team?

Either way, let me know. Thanks.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 3 - Message Variant 2',
        'subject': 'Re: reddit eating search volume',
        'body': '''Hey {FIRST_NAME}, we ran about 10 of these over the past few days. Pretty promising results.

What you get:
- 10 priority subreddits
- Strategy to capture 10-30k monthly reach from Reddit
- Exact posts to make, and where

I'm not certain if you're the best person for this though - is there someone else on your team I should send it to?

Appreciate your guidance either way.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 3 - Message Variant 3',
        'subject': 'Re: your reddit blind spot',
        'body': '''{FIRST_NAME} - we've done roughly 10 of these in the last week. Results have been impressive.

You'll walk away with:
- 10 priority subreddits to focus on
- How to generate 10-30k monthly visitors from Reddit
- Exactly what to post and where

Unsure if you're the one who should see this though. Should I connect with another teammate?

Let me know - thanks.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 3 - Message Variant 4',
        'subject': 'Re: 20% of search moved to reddit',
        'body': '''Hi {FIRST_NAME}, completed about 10 of these recently. Outcomes have been strong.

What you get:
- 10 priority subreddits
- Plan to achieve 10-30k monthly audience growth from Reddit
- Detailed posting strategy (what, where, how)

Not certain if this should go to you or someone else on your team though. Thoughts?

Thanks either way.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 3 - Message Variant 5',
        'subject': 'Re: reddits search takeover',
        'body': '''{FIRST_NAME}, we wrapped up about 10 of these this week. Results have been quite good.

Here's what you walk away with:
- 10 priority subreddits
- Path to 10-30k monthly visitors from Reddit
- Exact content strategy

Not sure if you're the right person though - would someone else on your team be better?

Let me know, thanks.

{SENDER_EMAIL_SIGNATURE}'''
    }
]

for msg in messages_step3_ceo:
    doc.add_heading(msg['title'], 2)
    doc.add_paragraph('Subject: ' + msg['subject'])
    doc.add_paragraph(msg['body'])

# Variant messages
messages_variant_ceo = [
    {
        'title': 'STEP 1 VARIANT - Message Variant 1',
        'subject': 'reddits overlooked value',
        'body': '''Hey {FIRST_NAME}, wild idea - want me to send you a free audit that outlines your current Reddit opportunities and how to capitalize so {HOOK} can hit the ground running?

Let me know, thanks.

{SENDER_EMAIL_SIGNATURE}

P.S. It typically uncovers 10-30k in monthly product-led search volume (the type where users show buying intent).'''
    },
    {
        'title': 'STEP 1 VARIANT - Message Variant 2',
        'subject': 'quick reddit opportunity',
        'body': '''{FIRST_NAME} - crazy thought. Would it be useful if I fired over a no-cost breakdown of how {COMPANY} can capitalize on Reddit opportunities?

Could help {HOOK} ramp up quickly.

Interested?

{SENDER_EMAIL_SIGNATURE}

P.S. Usually surfaces 10-30k monthly visitors with buying intent.'''
    },
    {
        'title': 'STEP 1 VARIANT - Message Variant 3',
        'subject': 'reddit growth angle',
        'body': '''Hi {FIRST_NAME},

Bold idea - I can send over a complimentary report covering your Reddit opportunities and exactly how to capitalize. Good onboarding tool for {HOOK}.

Cool if I share it?

{SENDER_EMAIL_SIGNATURE}

P.S. Typically reveals 10-30k monthly product-led volume (high buying intent).'''
    },
    {
        'title': 'STEP 1 VARIANT - Message Variant 4',
        'subject': 'reddit opportunities for {COMPANY}',
        'body': '''{FIRST_NAME},

Wild thought - would it be alright if I sent over a free audit covering {COMPANY}'s Reddit opportunities? Could help {HOOK} hit the ground running in the new role.

Worth a look?

{SENDER_EMAIL_SIGNATURE}

P.S. Usually uncovers 10-30k monthly visitors showing buying intent.'''
    },
    {
        'title': 'STEP 1 VARIANT - Message Variant 5',
        'subject': 'untapped reddit angle',
        'body': '''Hey {FIRST_NAME}, crazy idea - want me to shoot you a no-cost report that maps out your Reddit opportunities and how to capitalize?

Figure it'd help {HOOK} ramp up fast.

Lmk.

{SENDER_EMAIL_SIGNATURE}

P.S. Typically surfaces 10-30k in monthly product-led search (high buying intent).'''
    }
]

for msg in messages_variant_ceo:
    doc.add_heading(msg['title'], 2)
    doc.add_paragraph('Subject: ' + msg['subject'])
    doc.add_paragraph(msg['body'])

# CAMPAIGN 2
doc.add_page_break()
doc.add_heading('CAMPAIGN 2: New Hires - Marketing Leadership', 1)

# Step 1
messages_step1_marketing = [
    {
        'title': 'STEP 1 - Message Variant 1',
        'subject': 'reddits overlooked value',
        'body': '''Hey {FIRST_NAME}, saw you joined {COMPANY} as {TITLE} recently.

Would it help if I fire over a no-cost audit outlining your current opportunities on Reddit so you can hit the ground running?

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 1 - Message Variant 2',
        'subject': 'reddit opportunities',
        'body': '''{FIRST_NAME} - noticed you just started at {COMPANY} as {TITLE}.

Quick question: would it be valuable if I sent you a complimentary report that highlights your Reddit potential wins so you can ramp up quickly in your new role?

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 1 - Message Variant 3',
        'subject': 'new role reddit audit',
        'body': '''Hi {FIRST_NAME}, just noticed you joined {COMPANY} as {TITLE} pretty recently.

Would it be interesting if I shoot across a free deck mapping out your Reddit opportunities so you can hit the ground running?

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 1 - Message Variant 4',
        'subject': 'quick reddit win',
        'body': '''{FIRST_NAME},

Saw you recently joined {COMPANY} as {TITLE}.

Would it be worthwhile if I sent over an on-the-house audit that maps out your current Reddit potential so you can ramp up fast?

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 1 - Message Variant 5',
        'subject': 'reddit growth path',
        'body': '''Hey {FIRST_NAME}, noticed you joined {COMPANY} as {TITLE} not too long ago.

Quick ask - would it be cool if I fire over a complimentary report outlining your Reddit opportunities so you can hit the ground running in the new role?

{SENDER_EMAIL_SIGNATURE}'''
    }
]

for msg in messages_step1_marketing:
    doc.add_heading(msg['title'], 2)
    doc.add_paragraph('Subject: ' + msg['subject'])
    doc.add_paragraph(msg['body'])

# Step 1 Variant A
messages_step1_variantA_marketing = [
    {
        'title': 'STEP 1 VARIANT A - Message Variant 1',
        'subject': 'reddit opportunities',
        'body': '''{FIRST_NAME}, I figure ramping up quickly in your new role as {TITLE} at {COMPANY} is top of mind right now.

Quick question - would it be alright if I sent over a complimentary report covering your current opportunities on Reddit?

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 1 VARIANT A - Message Variant 2',
        'subject': 'ramping up at {COMPANY}',
        'body': '''{FIRST_NAME}, I imagine ramping up smoothly as {TITLE} at {COMPANY} is a big priority these days.

I wanted to ask - would it be okay if I shared a no-cost report highlighting your Reddit angles?

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 1 VARIANT A - Message Variant 3',
        'subject': 'new role support',
        'body': '''{FIRST_NAME}, I bet ramping up fast as {TITLE} at {COMPANY} is front and center atm.

Just wondering - would it be cool if I fired across a value-packed report that maps out your Reddit growth paths?

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 1 VARIANT A - Message Variant 4',
        'subject': 'reddit report',
        'body': '''{FIRST_NAME}, I imagine hitting the ground running in your new role as {TITLE} at {COMPANY} is top priority right now.

Quick question - would it be alright if I sent over an on-the-house report outlining your Reddit potential wins?

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 1 VARIANT A - Message Variant 5',
        'subject': 'onboarding fast at {COMPANY}',
        'body': '''{FIRST_NAME}, I figure ramping up fast as {TITLE} at {COMPANY} is a big priority these days.

I wanted to ask - would it be cool if I shared a complimentary report that covers your current Reddit opportunities?

{SENDER_EMAIL_SIGNATURE}'''
    }
]

for msg in messages_step1_variantA_marketing:
    doc.add_heading(msg['title'], 2)
    doc.add_paragraph('Subject: ' + msg['subject'])
    doc.add_paragraph(msg['body'])

# Step 1 Variant B
messages_step1_variantB_marketing = [
    {
        'title': 'STEP 1 VARIANT B - Message Variant 1',
        'subject': 'reddit opportunities',
        'body': '''Hey {FIRST_NAME}, would it be alright if I fired over a complimentary report covering your current Reddit opportunities so you can hit the ground running in your fresh role?

You'll see where to post, what to post, and how to get high-intent customers.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 1 VARIANT B - Message Variant 2',
        'subject': 'new role reddit playbook',
        'body': '''{FIRST_NAME}, would it be okay if I sent you a no-cost report outlining your Reddit opportunities so you can ramp up quickly in your new role?

Shows you where to post, what to post, how to convert high-intent visitors.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 1 VARIANT B - Message Variant 3',
        'subject': 'reddit growth report',
        'body': '''Hi {FIRST_NAME}, would it be cool if I shot you an on-the-house report that goes over your Reddit opportunities so you can hit the ground running?

You'll get: where to post, what content works, how to convert customers.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 1 VARIANT B - Message Variant 4',
        'subject': 'quick reddit win',
        'body': '''{FIRST_NAME}, would it be fine if I shared a value-packed report breaking down your Reddit opportunities so you can ramp up fast in your fresh role?

You'll see: target subreddits, content strategy, customer conversion tactics.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 1 VARIANT B - Message Variant 5',
        'subject': 'reddits hidden value',
        'body': '''Hey {FIRST_NAME}, would it be alright if I sent over a complimentary report that outlines your current Reddit opportunities so you can hit the ground running?

Covers: where to post, what to say, how to get high-intent customers.

{SENDER_EMAIL_SIGNATURE}'''
    }
]

for msg in messages_step1_variantB_marketing:
    doc.add_heading(msg['title'], 2)
    doc.add_paragraph('Subject: ' + msg['subject'])
    doc.add_paragraph(msg['body'])

# Step 2
messages_step2_marketing = [
    {
        'title': 'STEP 2 - Message Variant 1',
        'subject': 'Re: reddits overlooked value',
        'body': '''Reason I ask {FIRST_NAME} is because since the Google + Reddit partnership this year, Reddit now drives roughly 90% of high-intent searches.

This report will outline where your customers are spending time and how to engage with them so this becomes a measurable growth channel.

For example, when we worked with a personal finance app during the Mint shutdown:

6.2x growth in organic sessions
4.9x increase in installs from Reddit
5.1x lift from AI search mentions
37% reduction in support tickets

Nothing expected in return - would you like us to share a report that pinpoints 10-30k high-priority potential customers for {COMPANY}?

{SENDER_EMAIL_SIGNATURE}

P.S. We handle Reddit for some of the largest B2B SaaS in the world.'''
    },
    {
        'title': 'STEP 2 - Message Variant 2',
        'subject': 'Re: reddit opportunities',
        'body': '''The reason I'm reaching out {FIRST_NAME} is that after the Reddit-Google partnership earlier this year, Reddit drives about 90% of high-intent searches now.

The report will map out where your customers are active and how to connect with them so you get real business impact (not just vanity metrics).

Case in point: we partnered with a personal finance app during the Mint shutdown. They saw:

6.2x organic session growth
4.9x Reddit install increase
5.1x AI search lift
37% drop in support tickets

Just sharing in case it's useful - want us to put together a report surfacing 10,000-30,000 high-priority customers at {COMPANY}?

{SENDER_EMAIL_SIGNATURE}

P.S. We handle Reddit for major B2B SaaS companies.'''
    },
    {
        'title': 'STEP 2 - Message Variant 3',
        'subject': 'Re: new role reddit audit',
        'body': '''Let me explain why this matters {FIRST_NAME}. Since the Reddit-Google partnership, Reddit now drives nearly 90% of high-intent searches.

The report highlights where your customers are showing up and how to interact with them so this initiative drives more than vanity metrics.

Example: when we helped a personal finance app during Mint's shutdown, results were:

6.2x organic sessions
4.9x Reddit installs
5.1x AI search mentions
37% decrease in support tickets

Happy to pass this along - would you like us to provide a report mapping out ten to thirty thousand high-priority potential customers within {COMPANY}?

{SENDER_EMAIL_SIGNATURE}

P.S. We manage Reddit for some of the world's largest B2B SaaS.'''
    },
    {
        'title': 'STEP 2 - Message Variant 4',
        'subject': 'Re: quick reddit win',
        'body': '''Here's why I'm asking {FIRST_NAME}. After the Google + Reddit partnership earlier this year, Reddit drives about 90% of high-intent searches.

This report will show you where your customers are spending time and how to engage them to turn this into real business impact.

For reference, we worked with a personal finance app during the Mint shutdown and they experienced:

6.2x organic session growth
4.9x Reddit install lift
5.1x increase in AI search mentions
37% support ticket reduction

Nothing expected - would you like us to send over a report identifying 10-30k high-priority customers for {COMPANY}?

{SENDER_EMAIL_SIGNATURE}

P.S. We run Reddit for some of the biggest B2B SaaS in the world.'''
    },
    {
        'title': 'STEP 2 - Message Variant 5',
        'subject': 'Re: reddit growth path',
        'body': '''The reason I'm mentioning this {FIRST_NAME} is that since the Reddit-Google partnership, Reddit now drives roughly 90% of high-intent searches.

The report will outline where your customers are active and exactly how to engage them to make this a measurable growth channel.

To illustrate: we partnered with a personal finance app during Mint's shutdown. Results:

6.2x organic sessions
4.9x Reddit-driven installs
5.1x AI search lift
37% drop in support tickets

Just sharing in case it helps - want us to share a report that surfaces 10-30k high-priority potential customers at {COMPANY}?

{SENDER_EMAIL_SIGNATURE}

P.S. We handle Reddit for some of the largest B2B SaaS globally.'''
    }
]

for msg in messages_step2_marketing:
    doc.add_heading(msg['title'], 2)
    doc.add_paragraph('Subject: ' + msg['subject'])
    doc.add_paragraph(msg['body'])

# Step 3
messages_step3_marketing = [
    {
        'title': 'STEP 3 - Message Variant 1',
        'subject': 'Re: reddits overlooked value',
        'body': '''Looking back, maybe I've missed the mark {FIRST_NAME}. Are you the right person to review a report that uncovers overlooked customers on Reddit and outlines exactly how to capture them?

Let me know if I should speak with someone else.

Hope you crush it in the new role.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 3 - Message Variant 2',
        'subject': 'Re: reddit opportunities',
        'body': '''Thinking about it, maybe I've reached out to the wrong person {FIRST_NAME}. Are you the right contact to get a report that highlights untapped Reddit customers and details how to engage them?

Please let me know if I should connect with someone else.

Wishing you the best in the new role.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 3 - Message Variant 3',
        'subject': 'Re: new role reddit audit',
        'body': '''On reflection, maybe I've got this a little off {FIRST_NAME}. Are you the right person to see a report that surfaces under-the-radar customers on Reddit and shows exactly how to connect with them?

Happy to hear if I should reach out to someone else.

Best of luck as you ramp up at {COMPANY}.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 3 - Message Variant 4',
        'subject': 'Re: quick reddit win',
        'body': '''Thinking about it, I might have missed the mark {FIRST_NAME}. Are you the right contact to receive a report that identifies overlooked Reddit customers and outlines how to capture them?

Let me know if I should be talking to someone else.

Hope the new role goes well.

{SENDER_EMAIL_SIGNATURE}'''
    },
    {
        'title': 'STEP 3 - Message Variant 5',
        'subject': 'Re: reddit growth path',
        'body': '''Looking back, maybe I've reached out to the wrong person {FIRST_NAME}. Are you the right contact to get a report uncovering untapped Reddit customers and showing exactly how to engage them?

Please let me know if there's someone else I should connect with.

Wishing you success in your new role.

{SENDER_EMAIL_SIGNATURE}'''
    }
]

for msg in messages_step3_marketing:
    doc.add_heading(msg['title'], 2)
    doc.add_paragraph('Subject: ' + msg['subject'])
    doc.add_paragraph(msg['body'])

# Save
doc.save(r'c:\Users\mitch\Desktop\Claude Code Projects\Launclub_Anti_Fingerprinting_Variants.docx')
print("Document created successfully!")
