from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Launclub Email Sequence Variants - Anti-Fingerprinting', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Created using Josh Braun/Will Allred copywriting style')
doc.add_paragraph('5 variants per message to combat ESP fingerprinting')
doc.add_paragraph()

# Campaign 1: CEO/Founder/People Leaders
doc.add_heading('CAMPAIGN 1: New Hires - CEO, Founder, People Leaders', 1)

# Step 1 - Message 1
doc.add_heading('STEP 1 - Message Variant 1', 2)
doc.add_paragraph('Subject: reddit mayhem')
doc.add_paragraph(
    'Hey {FIRST_NAME}, saw Reddit basically hijacked 20% of Google search this year.\n\n'
    'Quick question - would it help if I sent over a breakdown of where your Reddit opportunities are '
    'and what to do about them?\n\n'
    'No pitch. Just figured it might be useful.\n\n'
    '{SENDER_EMAIL_SIGNATURE}\n\n'
    'P.S. We built Reddit strategies for Blizzard, Afterpay, and one of the Big 3 CRMs.'
)

doc.add_heading('STEP 1 - Message Variant 2', 2)
doc.add_paragraph('Subject: reddit eating search volume')
doc.add_paragraph(
    '{FIRST_NAME} - watching how Reddit absorbed a fifth of search traffic this year?\n\n'
    'Wondering if you\'d find value in a report showing exactly where {COMPANY} could be showing up '
    'on Reddit (and how to actually convert those visitors).\n\n'
    'Worth a look?\n\n'
    '{SENDER_EMAIL_SIGNATURE}\n\n'
    'P.S. We\'ve done this for brands like Blizzard and Afterpay - usually uncovers 10-30k monthly visitors.'
)

doc.add_heading('STEP 1 - Message Variant 3', 2)
doc.add_paragraph('Subject: your reddit blind spot')
doc.add_paragraph(
    'Hi {FIRST_NAME},\n\n'
    'Reddit now owns 20% of search. Most companies aren't adapting.\n\n'
    'I can send you a free audit showing your specific opportunities on Reddit - where to post, '
    'what to say, how to convert.\n\n'
    'Interested?\n\n'
    '{SENDER_EMAIL_SIGNATURE}\n\n'
    'P.S. We handle Reddit for Top 3 CRM and gaming companies like Blizzard.'
)

doc.add_heading('STEP 1 - Message Variant 4', 2)
doc.add_paragraph('Subject: 20% of search moved to reddit')
doc.add_paragraph(
    '{FIRST_NAME},\\n\\n'
    'Not sure if this is on your radar, but Reddit pulled in 20% of search volume this year.\\n\\n'
    'Would a report outlining {COMPANY}\\\'s current Reddit opportunities be useful? '
    'Shows you exactly which subreddits to target and what content converts.\\n\\n'
    'Free, obviously.\\n\\n'
    '{SENDER_EMAIL_SIGNATURE}\\n\\n'
    'P.S. Built Reddit programs for Afterpay, Blizzard, and a Top 3 CRM.'
)

doc.add_heading('STEP 1 - Message Variant 5', 2)
doc.add_paragraph('Subject: reddits search takeover')
doc.add_paragraph(
    'Hey {FIRST_NAME},\n\n'
    'Reddit absorbed 20% of search this year. Wild.\n\n'
    'Quick ask - want me to fire over a breakdown of your Reddit opportunities? '
    'It'll show you which subreddits matter and exactly what to post to generate conversions.\n\n'
    'Think it'd help.\n\n'
    '{SENDER_EMAIL_SIGNATURE}\n\n'
    'P.S. We've built Reddit strategies for orgs like Blizzard and Afterpay.'
)

# Step 2 - Follow up
doc.add_heading('STEP 2 - Message Variant 1', 2)
doc.add_paragraph('Subject: Re: reddit mayhem')
doc.add_paragraph(
    'Probably should\'ve been clearer {FIRST_NAME}.\n\n'
    'The report shows Reddit traffic opportunities based on searches you already rank for. '
    'Which subreddits to focus on. What posts to create. How to turn Reddit browsers into customers '
    '(and get picked up by LLMs while you\'re at it).\n\n'
    'If it\'s helpful, let me know and I\'ll get it over to you ASAP so {HOOK} can hit the ground running.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 2 - Message Variant 2', 2)
doc.add_paragraph('Subject: Re: reddit eating search volume')
doc.add_paragraph(
    '{FIRST_NAME} - realized I didn\'t explain this clearly.\n\n'
    'Basically, the report maps out Reddit opportunities based on keywords you currently rank for. '
    'You\'ll see:\n'
    '- Exact subreddits to target\n'
    '- Types of posts that convert\n'
    '- How to capture both Reddit traffic and LLM mentions\n\n'
    'Useful for {HOOK} to ramp up quickly in the new role. Want me to throw it together?\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 2 - Message Variant 3', 2)
doc.add_paragraph('Subject: Re: your reddit blind spot')
doc.add_paragraph(
    'Thinking about it, I could\'ve done a better job explaining this {FIRST_NAME}.\n\n'
    'The audit breaks down Reddit traffic opportunities tied to your current rankings. '
    'Shows you target subreddits, post types that drive conversions, and how to get traction in LLM results.\n\n'
    'If you think it\'d be useful for {HOOK} in the new role, give me a shout and I\'ll get it ready.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 2 - Message Variant 4', 2)
doc.add_paragraph('Subject: Re: 20% of search moved to reddit')
doc.add_paragraph(
    'On reflection, let me clarify {FIRST_NAME}.\n\n'
    'This report analyzes Reddit opportunities based on the searches you already rank on. '
    'It tells you which subreddits to post in, what content to create, and how to convert Reddit traffic '
    '(plus show up in LLM responses).\n\n'
    'If it sounds helpful for {HOOK} to hit the ground running, lmk and I\'ll put it together fast.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 2 - Message Variant 5', 2)
doc.add_paragraph('Subject: Re: reddits search takeover')
doc.add_paragraph(
    'Could\'ve been more specific {FIRST_NAME}.\n\n'
    'The breakdown shows Reddit opportunities based on keywords you currently rank for. '
    'You get: target subreddits, content to publish, conversion tactics, and LLM visibility strategies.\n\n'
    'Worth having for {HOOK} to ramp up quickly? Let me know and I\'ll send it over ASAP.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

# Step 3
doc.add_heading('STEP 3 - Message Variant 1', 2)
doc.add_paragraph('Subject: Re: reddit mayhem')
doc.add_paragraph(
    '{FIRST_NAME}, we\'ve completed about 10 of these this week. Results have been solid.\n\n'
    'You walk away with:\n'
    '- 10 priority subreddits to post in\n'
    '- Clear path to 10-30k monthly visitors from Reddit\n'
    '- Exactly what to post, where, and when\n\n'
    'Not sure if you\'re the right contact though. Should I be talking to someone else on your team?\n\n'
    'Either way, let me know. Thanks.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 3 - Message Variant 2', 2)
doc.add_paragraph('Subject: Re: reddit eating search volume')
doc.add_paragraph(
    'Hey {FIRST_NAME}, we ran about 10 of these over the past few days. Pretty promising results.\n\n'
    'What you get:\n'
    '- 10 priority subreddits\n'
    '- Strategy to capture 10-30k monthly reach from Reddit\n'
    '- Exact posts to make, and where\n\n'
    'I\'m not certain if you\'re the best person for this though - is there someone else on your team '
    'I should send it to?\n\n'
    'Appreciate your guidance either way.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 3 - Message Variant 3', 2)
doc.add_paragraph('Subject: Re: your reddit blind spot')
doc.add_paragraph(
    '{FIRST_NAME} - we\'ve done roughly 10 of these in the last week. Results have been impressive.\n\n'
    'You\'ll walk away with:\n'
    '- 10 priority subreddits to focus on\n'
    '- How to generate 10-30k monthly visitors from Reddit\n'
    '- Exactly what to post and where\n\n'
    'Unsure if you\'re the one who should see this though. Should I connect with another teammate?\n\n'
    'Let me know - thanks.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 3 - Message Variant 4', 2)
doc.add_paragraph('Subject: Re: 20% of search moved to reddit')
doc.add_paragraph(
    'Hi {FIRST_NAME}, completed about 10 of these recently. Outcomes have been strong.\n\n'
    'What you get:\n'
    '- 10 priority subreddits\n'
    '- Plan to achieve 10-30k monthly audience growth from Reddit\n'
    '- Detailed posting strategy (what, where, how)\n\n'
    'Not certain if this should go to you or someone else on your team though. Thoughts?\n\n'
    'Thanks either way.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 3 - Message Variant 5', 2)
doc.add_paragraph('Subject: Re: reddits search takeover')
doc.add_paragraph(
    '{FIRST_NAME}, we wrapped up about 10 of these this week. Results have been quite good.\n\n'
    'Here\'s what you walk away with:\n'
    '- 10 priority subreddits\n'
    '- Path to 10-30k monthly visitors from Reddit\n'
    '- Exact content strategy\n\n'
    'Not sure if you\'re the right person though - would someone else on your team be better?\n\n'
    'Let me know, thanks.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

# Variant Message
doc.add_heading('STEP 1 VARIANT - Message Variant 1', 2)
doc.add_paragraph('Subject: reddits overlooked value')
doc.add_paragraph(
    'Hey {FIRST_NAME}, wild idea - want me to send you a free audit that outlines your current Reddit '
    'opportunities and how to capitalize so {HOOK} can hit the ground running?\n\n'
    'Let me know, thanks.\n\n'
    '{SENDER_EMAIL_SIGNATURE}\n\n'
    'P.S. It typically uncovers 10-30k in monthly product-led search volume (the type where users show buying intent).'
)

doc.add_heading('STEP 1 VARIANT - Message Variant 2', 2)
doc.add_paragraph('Subject: quick reddit opportunity')
doc.add_paragraph(
    '{FIRST_NAME} - crazy thought. Would it be useful if I fired over a no-cost breakdown of how {COMPANY} '
    'can capitalize on Reddit opportunities?\n\n'
    'Could help {HOOK} ramp up quickly.\n\n'
    'Interested?\n\n'
    '{SENDER_EMAIL_SIGNATURE}\n\n'
    'P.S. Usually surfaces 10-30k monthly visitors with buying intent.'
)

doc.add_heading('STEP 1 VARIANT - Message Variant 3', 2)
doc.add_paragraph('Subject: reddit growth angle')
doc.add_paragraph(
    'Hi {FIRST_NAME},\n\n'
    'Bold idea - I can send over a complimentary report covering your Reddit opportunities and exactly '
    'how to capitalize. Good onboarding tool for {HOOK}.\n\n'
    'Cool if I share it?\n\n'
    '{SENDER_EMAIL_SIGNATURE}\n\n'
    'P.S. Typically reveals 10-30k monthly product-led volume (high buying intent).'
)

doc.add_heading('STEP 1 VARIANT - Message Variant 4', 2)
doc.add_paragraph('Subject: reddit opportunities for {COMPANY}')
doc.add_paragraph(
    '{FIRST_NAME},\n\n'
    'Wild thought - would it be alright if I sent over a free audit covering {COMPANY}\'s Reddit '
    'opportunities? Could help {HOOK} hit the ground running in the new role.\n\n'
    'Worth a look?\n\n'
    '{SENDER_EMAIL_SIGNATURE}\n\n'
    'P.S. Usually uncovers 10-30k monthly visitors showing buying intent.'
)

doc.add_heading('STEP 1 VARIANT - Message Variant 5', 2)
doc.add_paragraph('Subject: untapped reddit angle')
doc.add_paragraph(
    'Hey {FIRST_NAME}, crazy idea - want me to shoot you a no-cost report that maps out your Reddit '
    'opportunities and how to capitalize?\n\n'
    'Figure it\'d help {HOOK} ramp up fast.\n\n'
    'Lmk.\n\n'
    '{SENDER_EMAIL_SIGNATURE}\n\n'
    'P.S. Typically surfaces 10-30k in monthly product-led search (high buying intent).'
)

# CAMPAIGN 2
doc.add_page_break()
doc.add_heading('CAMPAIGN 2: New Hires - Marketing Leadership', 1)

# Step 1
doc.add_heading('STEP 1 - Message Variant 1', 2)
doc.add_paragraph('Subject: reddits overlooked value')
doc.add_paragraph(
    'Hey {FIRST_NAME}, saw you joined {COMPANY} as {TITLE} recently.\n\n'
    'Would it help if I fire over a no-cost audit outlining your current opportunities on Reddit '
    'so you can hit the ground running?\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 1 - Message Variant 2', 2)
doc.add_paragraph('Subject: reddit opportunities')
doc.add_paragraph(
    '{FIRST_NAME} - noticed you just started at {COMPANY} as {TITLE}.\n\n'
    'Quick question: would it be valuable if I sent you a complimentary report that highlights your '
    'Reddit potential wins so you can ramp up quickly in your new role?\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 1 - Message Variant 3', 2)
doc.add_paragraph('Subject: new role reddit audit')
doc.add_paragraph(
    'Hi {FIRST_NAME}, just noticed you joined {COMPANY} as {TITLE} pretty recently.\n\n'
    'Would it be interesting if I shoot across a free deck mapping out your Reddit opportunities '
    'so you can hit the ground running?\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 1 - Message Variant 4', 2)
doc.add_paragraph('Subject: quick reddit win')
doc.add_paragraph(
    '{FIRST_NAME},\n\n'
    'Saw you recently joined {COMPANY} as {TITLE}.\n\n'
    'Would it be worthwhile if I sent over an on-the-house audit that maps out your current Reddit '
    'potential so you can ramp up fast?\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 1 - Message Variant 5', 2)
doc.add_paragraph('Subject: reddit growth path')
doc.add_paragraph(
    'Hey {FIRST_NAME}, noticed you joined {COMPANY} as {TITLE} not too long ago.\n\n'
    'Quick ask - would it be cool if I fire over a complimentary report outlining your Reddit opportunities '
    'so you can hit the ground running in the new role?\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

# Step 1 Variants
doc.add_heading('STEP 1 VARIANT A - Message Variant 1', 2)
doc.add_paragraph('Subject: reddit opportunities')
doc.add_paragraph(
    '{FIRST_NAME}, I figure ramping up quickly in your new role as {TITLE} at {COMPANY} is top of mind right now.\n\n'
    'Quick question - would it be alright if I sent over a complimentary report covering your current '
    'opportunities on Reddit?\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 1 VARIANT A - Message Variant 2', 2)
doc.add_paragraph('Subject: ramping up at {COMPANY}')
doc.add_paragraph(
    '{FIRST_NAME}, I imagine ramping up smoothly as {TITLE} at {COMPANY} is a big priority these days.\n\n'
    'I wanted to ask - would it be okay if I shared a no-cost report highlighting your Reddit angles?\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 1 VARIANT A - Message Variant 3', 2)
doc.add_paragraph('Subject: new role support')
doc.add_paragraph(
    '{FIRST_NAME}, I bet ramping up fast as {TITLE} at {COMPANY} is front and center atm.\n\n'
    'Just wondering - would it be cool if I fired across a value-packed report that maps out your '
    'Reddit growth paths?\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 1 VARIANT A - Message Variant 4', 2)
doc.add_paragraph('Subject: reddit report')
doc.add_paragraph(
    '{FIRST_NAME}, I imagine hitting the ground running in your new role as {TITLE} at {COMPANY} is '
    'top priority right now.\n\n'
    'Quick question - would it be alright if I sent over an on-the-house report outlining your Reddit '
    'potential wins?\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 1 VARIANT A - Message Variant 5', 2)
doc.add_paragraph('Subject: onboarding fast at {COMPANY}')
doc.add_paragraph(
    '{FIRST_NAME}, I figure ramping up fast as {TITLE} at {COMPANY} is a big priority these days.\n\n'
    'I wanted to ask - would it be cool if I shared a complimentary report that covers your current '
    'Reddit opportunities?\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

# Step 1 Variant B
doc.add_heading('STEP 1 VARIANT B - Message Variant 1', 2)
doc.add_paragraph('Subject: reddit opportunities')
doc.add_paragraph(
    'Hey {FIRST_NAME}, would it be alright if I fired over a complimentary report covering your current '
    'Reddit opportunities so you can hit the ground running in your fresh role?\n\n'
    'You\'ll see where to post, what to post, and how to get high-intent customers.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 1 VARIANT B - Message Variant 2', 2)
doc.add_paragraph('Subject: new role reddit playbook')
doc.add_paragraph(
    '{FIRST_NAME}, would it be okay if I sent you a no-cost report outlining your Reddit opportunities '
    'so you can ramp up quickly in your new role?\n\n'
    'Shows you where to post, what to post, how to convert high-intent visitors.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 1 VARIANT B - Message Variant 3', 2)
doc.add_paragraph('Subject: reddit growth report')
doc.add_paragraph(
    'Hi {FIRST_NAME}, would it be cool if I shot you an on-the-house report that goes over your '
    'Reddit opportunities so you can hit the ground running?\n\n'
    'You\'ll get: where to post, what content works, how to convert customers.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 1 VARIANT B - Message Variant 4', 2)
doc.add_paragraph('Subject: quick reddit win')
doc.add_paragraph(
    '{FIRST_NAME}, would it be fine if I shared a value-packed report breaking down your Reddit '
    'opportunities so you can ramp up fast in your fresh role?\n\n'
    'You\'ll see: target subreddits, content strategy, customer conversion tactics.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 1 VARIANT B - Message Variant 5', 2)
doc.add_paragraph('Subject: reddits hidden value')
doc.add_paragraph(
    'Hey {FIRST_NAME}, would it be alright if I sent over a complimentary report that outlines your '
    'current Reddit opportunities so you can hit the ground running?\n\n'
    'Covers: where to post, what to say, how to get high-intent customers.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

# Step 2
doc.add_heading('STEP 2 - Message Variant 1', 2)
doc.add_paragraph('Subject: Re: reddits overlooked value')
doc.add_paragraph(
    'Reason I ask {FIRST_NAME} is because since the Google + Reddit partnership this year, '
    'Reddit now drives roughly 90% of high-intent searches.\n\n'
    'This report will outline where your customers are spending time and how to engage with them '
    'so this becomes a measurable growth channel.\n\n'
    'For example, when we worked with a personal finance app during the Mint shutdown:\n\n'
    '6.2x growth in organic sessions\n'
    '4.9x increase in installs from Reddit\n'
    '5.1x lift from AI search mentions\n'
    '37% reduction in support tickets\n\n'
    'Nothing expected in return - would you like us to share a report that pinpoints 10–30k high-priority '
    'potential customers for {COMPANY}?\n\n'
    '{SENDER_EMAIL_SIGNATURE}\n\n'
    'P.S. We handle Reddit for some of the largest B2B SaaS in the world.'
)

doc.add_heading('STEP 2 - Message Variant 2', 2)
doc.add_paragraph('Subject: Re: reddit opportunities')
doc.add_paragraph(
    'The reason I\'m reaching out {FIRST_NAME} is that after the Reddit–Google partnership earlier this year, '
    'Reddit drives about 90% of high-intent searches now.\n\n'
    'The report will map out where your customers are active and how to connect with them so you get '
    'real business impact (not just vanity metrics).\n\n'
    'Case in point: we partnered with a personal finance app during the Mint shutdown. They saw:\n\n'
    '6.2x organic session growth\n'
    '4.9x Reddit install increase\n'
    '5.1x AI search lift\n'
    '37% drop in support tickets\n\n'
    'Just sharing in case it\'s useful - want us to put together a report surfacing 10,000–30,000 '
    'high-priority customers at {COMPANY}?\n\n'
    '{SENDER_EMAIL_SIGNATURE}\n\n'
    'P.S. We handle Reddit for major B2B SaaS companies.'
)

doc.add_heading('STEP 2 - Message Variant 3', 2)
doc.add_paragraph('Subject: Re: new role reddit audit')
doc.add_paragraph(
    'Let me explain why this matters {FIRST_NAME}. Since the Reddit–Google partnership, Reddit now '
    'drives nearly 90% of high-intent searches.\n\n'
    'The report highlights where your customers are showing up and how to interact with them so this '
    'initiative drives more than vanity metrics.\n\n'
    'Example: when we helped a personal finance app during Mint\'s shutdown, results were:\n\n'
    '6.2x organic sessions\n'
    '4.9x Reddit installs\n'
    '5.1x AI search mentions\n'
    '37% decrease in support tickets\n\n'
    'Happy to pass this along - would you like us to provide a report mapping out ten to thirty thousand '
    'high-priority potential customers within {COMPANY}?\n\n'
    '{SENDER_EMAIL_SIGNATURE}\n\n'
    'P.S. We manage Reddit for some of the world\'s largest B2B SaaS.'
)

doc.add_heading('STEP 2 - Message Variant 4', 2)
doc.add_paragraph('Subject: Re: quick reddit win')
doc.add_paragraph(
    'Here\'s why I\'m asking {FIRST_NAME}. After the Google + Reddit partnership earlier this year, '
    'Reddit drives about 90% of high-intent searches.\n\n'
    'This report will show you where your customers are spending time and how to engage them to turn '
    'this into real business impact.\n\n'
    'For reference, we worked with a personal finance app during the Mint shutdown and they experienced:\n\n'
    '6.2x organic session growth\n'
    '4.9x Reddit install lift\n'
    '5.1x increase in AI search mentions\n'
    '37% support ticket reduction\n\n'
    'Nothing expected - would you like us to send over a report identifying 10–30k high-priority customers '
    'for {COMPANY}?\n\n'
    '{SENDER_EMAIL_SIGNATURE}\n\n'
    'P.S. We run Reddit for some of the biggest B2B SaaS in the world.'
)

doc.add_heading('STEP 2 - Message Variant 5', 2)
doc.add_paragraph('Subject: Re: reddit growth path')
doc.add_paragraph(
    'The reason I\'m mentioning this {FIRST_NAME} is that since the Reddit–Google partnership, '
    'Reddit now drives roughly 90% of high-intent searches.\n\n'
    'The report will outline where your customers are active and exactly how to engage them to make '
    'this a measurable growth channel.\n\n'
    'To illustrate: we partnered with a personal finance app during Mint\'s shutdown. Results:\n\n'
    '6.2x organic sessions\n'
    '4.9x Reddit-driven installs\n'
    '5.1x AI search lift\n'
    '37% drop in support tickets\n\n'
    'Just sharing in case it helps - want us to share a report that surfaces 10–30k high-priority '
    'potential customers at {COMPANY}?\n\n'
    '{SENDER_EMAIL_SIGNATURE}\n\n'
    'P.S. We handle Reddit for some of the largest B2B SaaS globally.'
)

# Step 3
doc.add_heading('STEP 3 - Message Variant 1', 2)
doc.add_paragraph('Subject: Re: reddits overlooked value')
doc.add_paragraph(
    'Looking back, maybe I\'ve missed the mark {FIRST_NAME}. Are you the right person to review a report '
    'that uncovers overlooked customers on Reddit and outlines exactly how to capture them?\n\n'
    'Let me know if I should speak with someone else.\n\n'
    'Hope you crush it in the new role.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 3 - Message Variant 2', 2)
doc.add_paragraph('Subject: Re: reddit opportunities')
doc.add_paragraph(
    'Thinking about it, maybe I\'ve reached out to the wrong person {FIRST_NAME}. Are you the right '
    'contact to get a report that highlights untapped Reddit customers and details how to engage them?\n\n'
    'Please let me know if I should connect with someone else.\n\n'
    'Wishing you the best in the new role.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 3 - Message Variant 3', 2)
doc.add_paragraph('Subject: Re: new role reddit audit')
doc.add_paragraph(
    'On reflection, maybe I\'ve got this a little off {FIRST_NAME}. Are you the right person to see '
    'a report that surfaces under-the-radar customers on Reddit and shows exactly how to connect with them?\n\n'
    'Happy to hear if I should reach out to someone else.\n\n'
    'Best of luck as you ramp up at {COMPANY}.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 3 - Message Variant 4', 2)
doc.add_paragraph('Subject: Re: quick reddit win')
doc.add_paragraph(
    'Thinking about it, I might have missed the mark {FIRST_NAME}. Are you the right contact to receive '
    'a report that identifies overlooked Reddit customers and outlines how to capture them?\n\n'
    'Let me know if I should be talking to someone else.\n\n'
    'Hope the new role goes well.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

doc.add_heading('STEP 3 - Message Variant 5', 2)
doc.add_paragraph('Subject: Re: reddit growth path')
doc.add_paragraph(
    'Looking back, maybe I\'ve reached out to the wrong person {FIRST_NAME}. Are you the right contact '
    'to get a report uncovering untapped Reddit customers and showing exactly how to engage them?\n\n'
    'Please let me know if there\'s someone else I should connect with.\n\n'
    'Wishing you success in your new role.\n\n'
    '{SENDER_EMAIL_SIGNATURE}'
)

# Save
doc.save(r'c:\Users\mitch\Desktop\Claude Code Projects\Launclub_Anti_Fingerprinting_Variants.docx')
print("Document created successfully!")
