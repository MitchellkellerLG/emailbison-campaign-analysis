from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Title
title = doc.add_heading('Cleanlab: AI Production Deployment Analysis & Messaging Strategy', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Comprehensive Market Research + Heavily Spintaxed Email Sequences')
doc.add_paragraph('Date: October 2025')
doc.add_paragraph()

# SECTION 1: MARKET RESEARCH
doc.add_heading('PART 1: MARKET RESEARCH & PATTERN INTERRUPTS', 1)

doc.add_heading('Enterprise AI Deployment Landscape 2025', 2)

doc.add_paragraph('Executive Summary:', style='Heading 3')
pain_points = [
    'Only 6% of enterprise GenAI projects reach production deployment (AWS Study)',
    '74% of enterprise CX AI initiatives fail due to poor data, strategy, and execution',
    'Traditional AI deployments take 12-18 months, while leading orgs now do it in 18 weeks',
    '39% of companies rank hallucinations as their #1 deployment challenge',
    'Hallucinations remain top concern for 38% of AI product leaders in customer-facing products',
    '50% of U.S. employees cite inaccuracy (including hallucinations) as top GenAI risk',
    'Stanford research found general LLMs hallucinated in 58-82% of legal queries',
    'Data annotation creates bottlenecks - enterprises need 70% time reduction through automation',
    'Only 31% of AI use cases reached full production in 2025 (doubled from 2024)',
    'CSAT benchmark is 80% for competitive industries, but 74% of chatbot programs fail',
]

for point in pain_points:
    p = doc.add_paragraph(point, style='List Bullet')

doc.add_heading('KEY PATTERN INTERRUPTS & ADDRESSABLE SITUATIONS', 2)

# Pattern Interrupt 1
doc.add_heading('Pattern Interrupt #1: "Artificial Stupidity" Crisis', 3)
doc.add_paragraph().add_run('Situation:').bold = True
doc.add_paragraph(
    'Companies are deploying AI that sounds intelligent but produces garbage outputs. '
    'ML teams spend months wrestling with hallucinations, retrieval misses, and noisy data. '
    'The traditional approach of manual spot-checking and labeling edge cases is a bottleneck.'
)

doc.add_paragraph().add_run('Why This Works:').bold = True
doc.add_paragraph(
    '• Frames AI failures as "stupidity" not "hallucinations" (more visceral)\n'
    '• Speaks to the exhaustion of ML teams stuck in 12-18 month deployment cycles\n'
    '• Pattern interrupt: "Stop wrestling with artificial stupidity"\n'
    '• Proof point: BBVA cut labeling time by 98% → weeks instead of months'
)

# Pattern Interrupt 2
doc.add_heading('Pattern Interrupt #2: The "CSAT Crashout" Problem', 3)
doc.add_paragraph().add_run('Situation:').bold = True
doc.add_paragraph(
    'Customer experience leaders are under pressure to deploy AI chatbots. They launch, '
    'CSAT drops, customers complain on reviews, team scrambles. 74% of CX AI programs fail. '
    'The automation-first approach (like Klarna replacing 700 employees) backfires when bots '
    'can\'t handle complex issues.'
)

doc.add_paragraph().add_run('Why This Works:').bold = True
doc.add_paragraph(
    '• "Crashout" is current slang that creates pattern interrupt\n'
    '• Speaks to fear of public failure (customers naming staff in bad reviews)\n'
    '• Positions 99% CSAT as the antidote to automation disasters\n'
    '• Social proof: Google, Alexa, BBVA achieved near-perfect outputs'
)

# Pattern Interrupt 3
doc.add_heading('Pattern Interrupt #3: The "Moral Obligation" Frame', 3)
doc.add_paragraph().add_run('Situation:').bold = True
doc.add_paragraph(
    'New hires in ML/AI roles are expected to deliver results fast. They see competitors '
    'achieving 99% accuracy in weeks. If Google, Alexa, and BBVA teams are using a system '
    'that works, wouldn\'t you be doing your company a disservice by not at least reviewing it?'
)

doc.add_paragraph().add_run('Why This Works:').bold = True
doc.add_paragraph(
    '• Reframes decision as ethical/professional duty, not sales pitch\n'
    '• Creates cognitive dissonance: "Am I ignoring proven solutions?"\n'
    '• Leverages social proof from tier-1 brands (Google, Alexa, BBVA)\n'
    '• Low commitment ask: "give it a review" not "buy it"'
)

# Pattern Interrupt 4
doc.add_heading('Pattern Interrupt #4: RAG Systems That "Sound Smart But Aren\'t Helpful"', 3)
doc.add_paragraph().add_run('Situation:').bold = True
doc.add_paragraph(
    'Enterprise teams build RAG systems that cite documents and sound authoritative, but '
    'answers are often irrelevant. The root cause: retrieval quality issues, outdated documents, '
    'and no systematic evaluation. Teams have "black box" pipelines that drift from accuracy over time.'
)

doc.add_paragraph().add_run('Why This Works:').bold = True
doc.add_paragraph(
    '• Identifies the gap between "sounds impressive" and "actually useful"\n'
    '• Addresses the evaluation/monitoring gap (most RAG deployed without metrics)\n'
    '• Pattern interrupt: "Your RAG is a black box slowly drifting away from accuracy"\n'
    '• Solution: Automatic flagging of hallucinations, retrieval misses, unreliable data'
)

# Pattern Interrupt 5
doc.add_heading('Pattern Interrupt #5: "Months to Weeks" Deployment Transformation', 3)
doc.add_paragraph().add_run('Situation:').bold = True
doc.add_paragraph(
    'Enterprises are stuck in 12-18 month AI deployment cycles while their competitors move '
    'from pilot to production in 18 weeks. The bottleneck: manual data labeling, spot-checking '
    'outputs, and chasing errors. AI-powered automation can reduce annotation time by 70%.'
)

doc.add_paragraph().add_run('Why This Works:').bold = True
doc.add_paragraph(
    '• Creates urgency: competitive disadvantage of slow deployment\n'
    '• Quantifies the transformation: 12-18 months → weeks\n'
    '• Addresses the real bottleneck: data labeling and quality control\n'
    '• Pattern interrupt: "Turn months of training into weeks"'
)

# WINNING MESSAGING THEMES
doc.add_heading('WINNING MESSAGING THEMES FOR CLEANLAB', 2)

themes = [
    ('Speed/Velocity', 'Months → Weeks, 98% labeling time reduction, "hit the ground running"'),
    ('Accuracy/Reliability', '99% CSAT, 98% accurate outputs, near-perfect customer experience'),
    ('Automation', 'Automatic flagging of hallucinations, zero infrastructure changes, no manual spot-checking'),
    ('Social Proof', 'Google, Amazon Alexa, BBVA, Fortune 1000 fintech'),
    ('Visceral Language', '"Artificial stupidity", "crashout", "wrestling with", "brawl with"'),
    ('Pattern Interrupts', '"Moral obligation", "weird one", "odd one", "interesting one"'),
    ('Low Friction', '"Brief email exchange", "quick look", "give it a review", "looksie"'),
    ('Situational Relevance', '"Fresh role", "new hire", "hit the ground running", "ramp up fast"'),
]

for theme, description in themes:
    p = doc.add_paragraph()
    run = p.add_run(theme + ': ')
    run.bold = True
    p.add_run(description)

# COMPANY SITUATIONS TO TARGET
doc.add_heading('COMPANY SITUATIONS TO TARGET', 2)

situations = [
    ('New AI/ML Hires', 'Recent joiners in ML, AI, CX, Customer Support roles - pressure to deliver fast wins'),
    ('CX Teams with Low CSAT', 'Companies with customer experience challenges, failed chatbot deployments'),
    ('RAG System Struggles', 'Teams deploying retrieval-augmented generation with accuracy issues'),
    ('Financial Services/Lending', 'Highly regulated industries where hallucinations = compliance risk'),
    ('Companies at AI Conferences', 'Attendees/engagers at AI events (AI4, Databricks posts) - active researchers'),
    ('Post-Series B/C Scale-ups', 'Companies 101-5000 employees trying to operationalize AI at scale'),
    ('Teams Using Generic LLMs', 'Organizations trying to use general-purpose models for specialized tasks'),
]

for situation, description in situations:
    p = doc.add_paragraph()
    run = p.add_run(situation + ': ')
    run.bold = True
    p.add_run(description)

# SECTION 2: SPINTAXED SEQUENCES
doc.add_page_break()
doc.add_heading('PART 2: HEAVILY SPINTAXED EMAIL SEQUENCES', 1)

doc.add_paragraph('Variables: {FIRST_NAME}, {SENDER_EMAIL_SIGNATURE}, {SENDER_FIRST_NAME} only')
doc.add_paragraph('Josh Braun/Will Allred style with anti-fingerprinting spintax')
doc.add_paragraph()

# CAMPAIGN 1: CSAT SENIOR LEADERSHIP
doc.add_heading('CAMPAIGN 1: CSAT - Senior Leadership', 1)

doc.add_heading('STEP 1 (Wait 2 Days)', 2)

csat_step1_v1 = {
    'subject': '{CX changes|customer experience|CSAT improvements}',
    'body': '''{Odd|Weird|Interesting} one {FIRST_NAME}, {noticed|saw|just saw} that you recently joined your {company|org|new company}.

{Wanted to ask|Quick question|Wondering} - if we could bring your CSAT to {99+%|near perfect|99% or higher} with {AI agents|chatbots|automated customer support} that {feels human|actually sounds human|doesn't feel robotic} ({without potential crashouts|without the typical AI failures|without embarrassing bot responses}), would a {quick|brief|short} email exchange be awful?

{SENDER_EMAIL_SIGNATURE}'''
}

csat_step1_v2 = {
    'subject': '{AI crashouts|chatbot failures|support automation disasters}',
    'body': '''{Odd|Weird|Interesting} one {FIRST_NAME}, the whole point of {artificial intelligence|AI|smart automation} in CX is {fast responses|quick resolution|instant answers} and {no crashouts|zero failures|no embarrassing mistakes} that get your staff named in reviews you might not like.

{But generalist models|Generic AI systems|Off-the-shelf models} make mistakes {(crashout)|(fail publicly)|(embarrass your team)} at an {even higher|alarmingly high|surprisingly high} rate.

We help teams build chatbots that {your grandma would even enjoy|customers actually love|feel genuinely helpful} with {99+%|near-perfect|99% or higher} CSAT in {weeks|a few weeks|just weeks}.

{Open to seeing a few of them in action?|Worth a quick look?|Interested in seeing how?}

{SENDER_EMAIL_SIGNATURE}'''
}

csat_step1_v3 = {
    'subject': '{moral obligation|professional duty|quick question}',
    'body': '''{FIRST_NAME}, if {Google|Amazon Alexa|BBVA|Google and Alexa|Google, Alexa and BBVA} teams all achieved {near perfect|99+%|99% or higher} CSAT on their custom chat bots would you feel morally {opposed|obligated|compelled} to give the {process|system|approach} they used a {looksie|review|run through}?

{SENDER_FIRST_NAME}

P.S. {thought it might be relevant given the new role|figured it'd be useful as you ramp up|seemed relevant for your new position}.'''
}

doc.add_heading('VARIANT 1', 3)
doc.add_paragraph('Subject: ' + csat_step1_v1['subject'])
doc.add_paragraph(csat_step1_v1['body'])

doc.add_heading('VARIANT 2', 3)
doc.add_paragraph('Subject: ' + csat_step1_v2['subject'])
doc.add_paragraph(csat_step1_v2['body'])

doc.add_heading('VARIANT 3', 3)
doc.add_paragraph('Subject: ' + csat_step1_v3['subject'])
doc.add_paragraph(csat_step1_v3['body'])

# Step 2
doc.add_heading('STEP 2 (Wait 5 Days After Step 1)', 2)

csat_step2_v1 = {
    'subject': 'Re: {CX changes|customer experience|CSAT improvements}',
    'body': '''You know what {FIRST_NAME}, I {completely|totally|entirely} forgot to {share|mention|explain} how we're helping companies {build|create|deploy} AI chatbots that customers {love|actually trust|genuinely appreciate}.

Picture this:
- You're running AI-powered {support|service|customer care} channels
- Instead of manually {reviewing|spot-checking|verifying} {responses|answers|outputs} or chasing errors, Cleanlab automatically {flags|detects|catches} hallucinations, retrieval misses, and {unreliable|messy|noisy} data
- Suddenly {responses|answers|interactions} start to feel human, and your team is {alerted|informed|notified} exactly when they need to step in

Right now, your AI probably {routes|sends|escalates} some {stuff|issues|tickets} to agents - but you're not sure when it should.

Cleanlab figures that out. It {catches|identifies|flags} the {risky|problematic|dangerous} conversations before they go sideways and {routes|sends|escalates} them to your team.

Think {fewer|less|minimal} {angry|frustrated|upset} escalations. More customers who actually {trust|appreciate|value} your support.

If you'd like, we can {schedule|set up|arrange} a look at your current AI support system and see how Cleanlab could help {improve|boost|elevate} your CSAT and CX {metrics|scores|performance}.

{SENDER_EMAIL_SIGNATURE}

{Btw|By the way|P.S.} this is why teams at {Google, Alexa and BBVA|Google, Amazon Alexa, and BBVA|major brands like Google and BBVA} trust us to train their systems.'''
}

csat_step2_v2 = {
    'subject': 'Re: {AI crashouts|chatbot failures|support automation disasters}',
    'body': '''{Looking back|Thinking about it|On reflection} {FIRST_NAME}, I {should've explained better|forgot to mention the details|could've been clearer}.

{Here's how it works|The approach is simple|Here's the system}:

Your team is {running|deploying|managing} AI-powered {support|customer service|care} channels. {Instead of|Rather than} manually {reviewing|checking|verifying} every {response|answer|interaction} or {hunting down|chasing|tracking} errors, Cleanlab {automatically|instantly} {flags|detects|identifies}:
- Hallucinations
- Retrieval misses
- {Unreliable|Noisy|Messy} data

{The result?|What happens?|The outcome?} {Responses|Answers|Interactions} {start to|begin to|suddenly} feel human. Your team gets {alerted|notified|pinged} exactly when they need to {step in|intervene|take over}.

Think {fewer|less|minimal} {angry|frustrated|escalated} customer {issues|complaints|tickets}. {More|Better|Higher} {trust|satisfaction|appreciation} in your support.

{Would a brief look at your current setup make sense?|Want to see how this could work for you?|Open to a quick review of your system?}

{SENDER_EMAIL_SIGNATURE}

{Btw|P.S.|By the way} - teams at {Google and BBVA|Amazon Alexa and Google|Fortune 1000 fintech companies} rely on our work.'''
}

csat_step2_v3 = {
    'subject': 'Re: {moral obligation|professional duty|quick question}',
    'body': '''{FIRST_NAME}, {let me explain why I asked|here's why this matters|quick context}:

{Google|Amazon Alexa|BBVA|Major enterprises} {achieved|reached|hit} {99+%|near-perfect|99% or higher} CSAT with {AI chatbots|automated support|AI agents} in {weeks|just weeks|a few weeks}, not months.

The {difference|key|secret}? They {stopped|quit|avoided} manually {reviewing|spot-checking|verifying} every {response|output|answer}.

Cleanlab {automatically|instantly} {flags|catches|detects} hallucinations, retrieval {misses|failures|gaps}, and {unreliable|messy|noisy} data - so your team {knows|gets alerted|is notified} exactly when to {step in|intervene|take over}.

{The result|What happens|The outcome}: {Fewer|Less|Minimal} {angry|frustrated} escalations. {More|Better|Higher} customer {trust|satisfaction|appreciation}.

{Would a quick look at your current system make sense?|Open to seeing how this works?|Worth a brief review?}

{SENDER_EMAIL_SIGNATURE}'''
}

doc.add_heading('VARIANT 1', 3)
doc.add_paragraph('Subject: ' + csat_step2_v1['subject'])
doc.add_paragraph(csat_step2_v1['body'])

doc.add_heading('VARIANT 2', 3)
doc.add_paragraph('Subject: ' + csat_step2_v2['subject'])
doc.add_paragraph(csat_step2_v2['body'])

doc.add_heading('VARIANT 3', 3)
doc.add_paragraph('Subject: ' + csat_step2_v3['subject'])
doc.add_paragraph(csat_step2_v3['body'])

# Step 3
doc.add_heading('STEP 3 (Wait 1 Day After Step 2)', 2)

csat_step3_v1 = {
    'subject': 'Re: {CX changes|customer experience|CSAT improvements}',
    'body': '''Hi {FIRST_NAME}, I get that with the scope of your team, I might have {missed|reached out to} the {right|wrong} person to discuss {improving|boosting|elevating} your {CSAT|CX|customer experience scores}.

Would someone else in the team be the right person to connect with about this?

Alternatively, I can {send|share|forward} a report on how we perform across key LLM {benchmark|metrics|indicators}.

{SENDER_EMAIL_SIGNATURE}

{Btw|P.S.|By the way} the cool thing about our process is that we actually published {award winning|award-winning|recognized} research on AI training and our proprietary software {based on that|from that research|leveraging that} means that teams at {Google, Alexa, and fortune 1000 fintech companies|Google, Amazon Alexa, and BBVA|major enterprises like Google and BBVA} rely on our work.'''
}

csat_step3_v2 = {
    'subject': 'Re: {AI crashouts|chatbot failures|support automation disasters}',
    'body': '''{FIRST_NAME}, {thinking about it|on reflection|looking back}, maybe I {reached out to the wrong person|missed the mark|should've connected with someone else}.

Are you the right {person|contact|teammate} to {discuss|review|explore} {improving|boosting|elevating} {CSAT|customer experience|CX scores} with {AI automation|automated support|AI agents}?

{Or|Otherwise|Alternatively} {should I connect with someone else|is there a better person|would another teammate be better}?

{Happy to send|Can share|Can forward} a report on our {LLM benchmarks|performance metrics|accuracy indicators} either way.

{SENDER_EMAIL_SIGNATURE}'''
}

csat_step3_v3 = {
    'subject': 'Re: {moral obligation|professional duty|quick question}',
    'body': '''{Hey|Hi} {FIRST_NAME}, I {might've|may have|could've} {reached out to|contacted} the wrong person.

Are you the {right|best} {contact|person|teammate} to {see|review|get} a report on how {we help companies achieve|teams like Google and BBVA achieved|enterprises reach} {99+% CSAT|near-perfect customer satisfaction|99% or higher CSAT} with AI {chatbots|support|agents}?

{Or|Otherwise|If not} {should I be talking to|would you recommend|is there} someone else {on your team|at the company}?

{SENDER_EMAIL_SIGNATURE}

P.S. We published {award-winning|recognized|peer-reviewed} research on AI training that {Google, Alexa, and Fortune 1000 companies|major enterprises|top-tier brands} use.'''
}

doc.add_heading('VARIANT 1', 3)
doc.add_paragraph('Subject: ' + csat_step3_v1['subject'])
doc.add_paragraph(csat_step3_v1['body'])

doc.add_heading('VARIANT 2', 3)
doc.add_paragraph('Subject: ' + csat_step3_v2['subject'])
doc.add_paragraph(csat_step3_v2['body'])

doc.add_heading('VARIANT 3', 3)
doc.add_paragraph('Subject: ' + csat_step3_v3['subject'])
doc.add_paragraph(csat_step3_v3['body'])

# CAMPAIGN 2: RECENT AI HIRES
doc.add_page_break()
doc.add_heading('CAMPAIGN 2: Recent AI/ML Hires', 1)

doc.add_heading('STEP 1 (Wait 1 Day)', 2)

ai_step1_v1 = {
    'subject': '{artificial stupidity|AI failures|model accuracy issues}',
    'body': '''{Hey|Hi|Hello} {FIRST_NAME},

We're helping dozens of ML teams {stop|avoid|eliminate} wrestling with {AI stupidity|artificial stupidity|model failures} by turning {months|months and months} of training into {weeks|just weeks|a few weeks} ({98%|98% or higher|near-perfect} {accurate|reliable|precise} outputs).

Anything you're {trying|looking|planning} to get into production {atm|right now|currently}?

{SENDER_EMAIL_SIGNATURE}'''
}

ai_step1_v2 = {
    'subject': '{production AI challenges|getting models live|deployment struggles}',
    'body': '''{Hey|Hi} {FIRST_NAME}, {odd|weird|interesting} question:

Are you {trying|pushing|fighting} to get {any AI|ML models|custom AI systems} into production {right now|atm|currently}?

We're seeing ML teams {cut|reduce|slash} training time from {months to weeks|12-18 months to just weeks|months down to weeks} with {98%|near-perfect|98% or higher} accuracy.

{Worth a quick look?|Make sense to explore?|Interested?}

{SENDER_EMAIL_SIGNATURE}'''
}

ai_step1_v3 = {
    'subject': '{morally opposed|professional obligation|quick question}',
    'body': '''{FIRST_NAME}, if {Google|Amazon Alexa|BBVA|Google and Alexa|Google, Alexa and BBVA} teams all achieved {near perfect|99+%|99% or higher} outputs on their custom {chat bots|AI models|ML systems} in {weeks|just weeks|a few weeks} would you feel morally {opposed|obligated|compelled} to give the {process|system|approach} they used a {looksie|review|run through}?

{SENDER_FIRST_NAME}

P.S. {thought it might be interesting|figured it'd be relevant|seemed worth mentioning} given the recent ML team {additions|hires|expansion}.'''
}

doc.add_heading('VARIANT 1', 3)
doc.add_paragraph('Subject: ' + ai_step1_v1['subject'])
doc.add_paragraph(ai_step1_v1['body'])

doc.add_heading('VARIANT 2', 3)
doc.add_paragraph('Subject: ' + ai_step1_v2['subject'])
doc.add_paragraph(ai_step1_v2['body'])

doc.add_heading('VARIANT 3', 3)
doc.add_paragraph('Subject: ' + ai_step1_v3['subject'])
doc.add_paragraph(ai_step1_v3['body'])

# AI Step 2
doc.add_heading('STEP 2 (Wait 5 Days After Step 1)', 2)

ai_step2_v1 = {
    'subject': 'Re: {artificial stupidity|AI failures|model accuracy issues}',
    'body': '''{Looking back|Thinking about it|On reflection} {FIRST_NAME}, I {completely|totally|entirely} forgot to share how we're solving the {artificial stupidity|AI accuracy|model reliability} crisis.

Picture this:
- You're in a {brawl|fight|struggle} with a new {RAG|AI|ML} {agent|assistant|system}
- Instead of manually {spot-checking|reviewing|verifying} outputs or labeling edge cases, Cleanlab {flags|detects|catches} hallucinations, retrieval misses, and {noisy|messy|unreliable} data {automatically|instantly}
- Your model {gets to play nice with|starts working properly for|becomes reliable with} customers in {weeks|just weeks|a few weeks} {instead of|rather than|not} months

That's exactly what happened at BBVA. They cut labeling time by {98%|nearly all of it|almost entirely} after {plugging|integrating|connecting} Cleanlab into their ML workflows ({with zero infrastructure changes|without changing infrastructure|no infrastructure overhaul}).

It takes any AI from {spewing|spouting|producing} nonsense to {speaking|communicating|responding} properly in {weeks|just weeks}.

Would a {brief|quick|short} email exchange around model training be awful?

{SENDER_EMAIL_SIGNATURE}

P.S. Most of our clients reach {98%|near-perfect|98% or higher} accuracy on their AI workflows.'''
}

ai_step2_v2 = {
    'subject': 'Re: {production AI challenges|getting models live|deployment struggles}',
    'body': '''{FIRST_NAME}, I {should've explained this better|forgot to mention how it works|could've been clearer}.

{Here's the situation|The problem we solve|What we do}:

You're {building|deploying|training} a {RAG system|custom AI model|ML agent}. {Instead of|Rather than} manually {spot-checking|reviewing|labeling} every {output|response|edge case}, Cleanlab {automatically|instantly} {flags|detects|identifies}:
- Hallucinations
- Retrieval {misses|failures|gaps}
- {Noisy|Messy|Unreliable} data

{The result?|What happens?|The outcome?} Your model {goes from|shifts from|moves from} {producing|generating|spewing} nonsense to {working|performing|responding} properly in {weeks|just weeks|a few weeks}, not months.

{Example|Case study|Real world}: BBVA cut labeling time by {98%|nearly everything|almost entirely}. {Zero|No|Without any} infrastructure changes.

{Would a brief exchange make sense?|Worth exploring?|Open to a quick look?}

{SENDER_EMAIL_SIGNATURE}'''
}

ai_step2_v3 = {
    'subject': 'Re: {morally opposed|professional obligation|quick question}',
    'body': '''{FIRST_NAME}, {here's why I asked|let me explain|quick context}:

{Google|Amazon Alexa|BBVA|Major ML teams} {achieved|reached|hit} {98%|near-perfect|98% or higher} accuracy in {weeks|just weeks|a few weeks} by {stopping|avoiding|eliminating} manual {spot-checking|output review|edge case labeling}.

Cleanlab {automatically|instantly} {flags|catches|detects} hallucinations, retrieval {misses|failures|gaps}, and {unreliable|noisy|messy} data.

{The result|What happens|The outcome}: Models go from {producing|generating|spewing} nonsense to {working|responding|performing} properly in {weeks|just weeks}, not months.

{BBVA|One fintech client|A Fortune 1000 company} cut labeling time by {98%|nearly all of it}.

{Would a quick look make sense?|Worth exploring?|Open to a brief exchange?}

{SENDER_EMAIL_SIGNATURE}'''
}

doc.add_heading('VARIANT 1', 3)
doc.add_paragraph('Subject: ' + ai_step2_v1['subject'])
doc.add_paragraph(ai_step2_v1['body'])

doc.add_heading('VARIANT 2', 3)
doc.add_paragraph('Subject: ' + ai_step2_v2['subject'])
doc.add_paragraph(ai_step2_v2['body'])

doc.add_heading('VARIANT 3', 3)
doc.add_paragraph('Subject: ' + ai_step2_v3['subject'])
doc.add_paragraph(ai_step2_v3['body'])

# AI Step 3
doc.add_heading('STEP 3 (Wait 1 Day After Step 2)', 2)

ai_step3_v1 = {
    'subject': 'Re: {artificial stupidity|AI failures|model accuracy issues}',
    'body': '''{Btw|By the way|Quick note} {FIRST_NAME}, worth noting we're seeing {nearly perfect|99+%|near-perfect} customer experience scores with production LLM systems for companies like {Google, Amazon Alexa and BBVA|Google and BBVA|major enterprises}.

{SENDER_FIRST_NAME}

Actually did some benchmarking that {lays out|outlines|shows} the best method for building these out in {weeks|just weeks} - {want the report?|interested in seeing it?|worth a look?} or maybe {someone|somebody} else on the team might be interested?'''
}

ai_step3_v2 = {
    'subject': 'Re: {production AI challenges|getting models live|deployment struggles}',
    'body': '''{FIRST_NAME}, {thinking about it|on reflection|looking back}, maybe I {reached out to the wrong person|should've connected with someone else|missed the right contact}.

Are you the {right|best} person to {see|review|get} a report on how {we help ML teams|teams like Google and BBVA|enterprises} {achieve|reach|hit} {98%|near-perfect|98% or higher} accuracy in {weeks|just weeks}?

{Or|Otherwise|If not} {should I be talking to|would you recommend|is there} someone else {on your team|at the company}?

{SENDER_EMAIL_SIGNATURE}'''
}

ai_step3_v3 = {
    'subject': 'Re: {morally opposed|professional obligation|quick question}',
    'body': '''{Hey|Hi} {FIRST_NAME}, I {might've|may have|could've} {reached out to|contacted} the wrong person.

Are you the {right|best} {contact|person|teammate} to {see|review|discuss} how {Google, Alexa and BBVA|major ML teams|enterprises} {achieved|reached|got to} {98%|near-perfect} accuracy in {weeks|just weeks}?

{Or|If not|Otherwise} {should I connect with|would you recommend|is there} someone else?

{SENDER_FIRST_NAME}

P.S. We have benchmarking {data|reports|research} that {shows|outlines|details} the {process|approach|method}.'''
}

doc.add_heading('VARIANT 1', 3)
doc.add_paragraph('Subject: ' + ai_step3_v1['subject'])
doc.add_paragraph(ai_step3_v1['body'])

doc.add_heading('VARIANT 2', 3)
doc.add_paragraph('Subject: ' + ai_step3_v2['subject'])
doc.add_paragraph(ai_step3_v2['body'])

doc.add_heading('VARIANT 3', 3)
doc.add_paragraph('Subject: ' + ai_step3_v3['subject'])
doc.add_paragraph(ai_step3_v3['body'])

# IMPLEMENTATION NOTES
doc.add_page_break()
doc.add_heading('IMPLEMENTATION NOTES', 1)

notes = [
    ('Variables Used', 'Only {FIRST_NAME}, {SENDER_EMAIL_SIGNATURE}, {SENDER_FIRST_NAME} - removed all {COMPANY}, {TITLE}, {HOOK} variables'),
    ('Spintax Density', 'Heavy spintax on greetings, verbs, value props, social proof, CTAs, and technical terms'),
    ('Spam Avoidance', 'No excessive caps, multiple !!!, "free" overuse, urgency triggers, or promotional language'),
    ('Pattern Interrupts', '"Artificial stupidity", "crashout", "moral obligation", "odd/weird one"'),
    ('Social Proof', 'Google, Amazon Alexa, BBVA, Fortune 1000 - tier 1 brand validation'),
    ('Quantified Benefits', '98% labeling reduction, 99% CSAT, weeks not months, zero infrastructure changes'),
    ('Josh Braun Style', 'Short sentences, visceral language, "brawl with", "wrestling with", casual CTAs'),
    ('Will Allred Style', 'Direct value, P.S. social proof, low-friction asks, situational relevance'),
    ('Thread Continuity', 'Step 2 and 3 use "Re:" with same subject spintax to maintain threading'),
    ('Wait Times', 'CSAT: 2 days → 5 days → 1 day | AI Hires: 1 day → 5 days → 1 day'),
]

for note_title, note_content in notes:
    p = doc.add_paragraph()
    run = p.add_run(note_title + ': ')
    run.bold = True
    p.add_run(note_content)

# TARGETING RECOMMENDATIONS
doc.add_heading('TARGETING RECOMMENDATIONS', 2)

targeting = [
    ('Campaign 1 (CSAT)', 'Target: VP/Director of Customer Experience, Support, CX roles at B2C/B2B SaaS companies with customer-facing AI'),
    ('Campaign 2 (AI Hires)', 'Target: ML Engineers, AI/ML leads, Data Scientists hired in last 90 days at companies deploying production AI'),
    ('Industry Focus', 'Financial Services (lending, credit), SaaS, E-commerce, FinTech - high compliance + accuracy needs'),
    ('Company Size', '101-5000 employees (Series B/C+ stage) - past experimentation, struggling with production deployment'),
    ('Trigger Events', 'New hires, AI conference attendees, low CSAT scores, failed chatbot deployments, RAG system launches'),
    ('Timing', 'Send CSAT campaign after Q1/Q3 CSAT reviews when scores are discussed; AI hire campaign within 30-60 days of hire'),
]

for target_title, target_content in targeting:
    p = doc.add_paragraph()
    run = p.add_run(target_title + ': ')
    run.bold = True
    p.add_run(target_content)

# Save
doc.save(r'c:\Users\mitch\Desktop\Claude Code Projects\Cleanlab_Comprehensive_Analysis.docx')
print("Comprehensive Cleanlab analysis document created successfully!")
